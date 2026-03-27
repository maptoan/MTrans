# -*- coding: utf-8 -*-

"""
OCR reader module: extract text from scanned PDFs or images based on settings in config/config.yaml.

Dependencies:
- pytesseract (Python wrapper for Tesseract OCR)
- pdf2image (convert PDF pages to images)
- Pillow (image processing)
- PyYAML (read YAML config)

Config example in config/config.yaml:
  ocr:
    enabled: true
    tesseract_cmd: "C:/Program Files/Tesseract-OCR/tesseract.exe"
    lang: "vie+eng"
    psm: 3
    dpi: 300
"""

import asyncio
import gc
import logging
import os
import subprocess
import sys
import time
from pathlib import Path
from typing import List, Optional

# Suppress noisy logs từ Google libraries (absl, gRPC) trước khi import
os.environ["GRPC_VERBOSITY"] = "ERROR"
os.environ["GLOG_minloglevel"] = "2"  # Suppress INFO và WARNING từ GLOG/absl
os.environ["GRPC_PYTHON_LOG_LEVEL"] = "ERROR"
os.environ["TF_CPP_MIN_LOG_LEVEL"] = "2"  # Suppress TensorFlow logs nếu có


# Define StderrFilter trước để filter stderr ngay từ đầu
class NoisyMessageFilter:
    """Filter để chặn các messages gây nhiễu được in trực tiếp ra stderr/stdout."""

    def __init__(self, original_stream):
        self.original_stream = original_stream
        # Buffer để xử lý multi-line messages
        self.buffer = ""
        # Patterns gây nhiễu (tất cả lowercase để so sánh)
        # Bao gồm cả partial matches để catch variations
        self.noisy_patterns = [
            "e0000",  # gRPC error prefix
            "alts_credentials",
            "alts creds",
            "alts creds ignored",
            "alts creds ignored. not running on gcp",
            "absl::initializelog",
            "not running on gcp",
            "untrusted alts",
            "untrusted alts is not enabled",
            "written to stderr",
            "all log messages before",
            "all log messages before absl",
            "alts_credentials.cc",  # File path pattern
            "alts_credentials.cc:93",  # File path with line number
            "warning: all log messages",  # Warning prefix
        ]

    def write(self, text):
        if not text:
            return

        # Thêm vào buffer để xử lý multi-line messages
        self.buffer += text

        # Kiểm tra buffer có chứa noisy patterns không (case-insensitive)
        buffer_lower = self.buffer.lower()

        # Kiểm tra nhanh trước khi split (tối ưu hơn)
        is_noisy = False

        # Check toàn bộ buffer trước (faster)
        for pattern in self.noisy_patterns:
            if pattern in buffer_lower:
                is_noisy = True
                break

        # Nếu chưa detect, check từng dòng chi tiết
        if not is_noisy:
            lines = self.buffer.split("\n")
            for line in lines:
                line_lower = line.lower().strip()
                # Check các pattern cụ thể
                if any(pattern in line_lower for pattern in self.noisy_patterns):
                    is_noisy = True
                    break
                # Check pattern E0000 ở đầu dòng
                if line_lower.startswith("e0000"):
                    is_noisy = True
                    break
                # Check "WARNING:" prefix với absl messages
                if line_lower.startswith("warning:") and (
                    "absl" in line_lower or "stderr" in line_lower
                ):
                    is_noisy = True
                    break

        # Nếu không noisy, ghi ra stream
        if not is_noisy:
            self.original_stream.write(text)
        # Nếu noisy, không ghi gì cả (suppress hoàn toàn)

        # Reset buffer sau mỗi newline hoặc khi buffer quá dài
        if "\n" in text:
            # Giữ lại phần sau newline cuối cùng để check tiếp (cho multi-line messages)
            parts = self.buffer.rsplit("\n", 1)
            self.buffer = parts[-1] if len(parts) > 1 else ""

        if len(self.buffer) > 2000:  # Reset nếu buffer quá dài
            self.buffer = ""

    def flush(self):
        self.original_stream.flush()

    def __getattr__(self, name):
        return getattr(self.original_stream, name)


# Alias cho backward compatibility
StderrFilter = NoisyMessageFilter

# Suppress warnings từ absl trước khi import google libraries
try:
    import absl.logging

    absl.logging.set_verbosity(absl.logging.ERROR)
except Exception:
    pass

# Suppress logging từ các Google libraries
for lib_name in [
    "google",
    "grpc",
    "absl",
    "google.generativeai",
    "google.api_core",
    "google.auth",
]:
    lib_logger = logging.getLogger(lib_name)
    lib_logger.setLevel(logging.ERROR)
    lib_logger.propagate = False

# Filter stderr ngay từ đầu để chặn messages in trực tiếp
# (không filter stdout ở đây để không ảnh hưởng đến user interaction)
_stderr_filter_active = False
_stdout_filter_active = False
try:
    original_stderr = sys.stderr
    sys.stderr = NoisyMessageFilter(original_stderr)
    _stderr_filter_active = True
except Exception:
    pass

import yaml

try:
    from PIL import Image
except Exception:
    Image = None

try:
    import pytesseract
except Exception:
    pytesseract = None

try:
    from pdf2image import convert_from_path
except Exception:
    convert_from_path = None

try:
    from tqdm import tqdm  # progress bar (optional)
except Exception:
    tqdm = None

try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from docx import Document
    from docx.shared import Inches, Pt
except Exception:
    Document = None

logger = logging.getLogger("NovelTranslator")


class GoogleLogFilter(logging.Filter):
    """
    Filter để loại bỏ các log messages gây nhiễu từ Google libraries.
    """

    def filter(self, record):
        msg = str(record.getMessage())
        msg_lower = msg.lower()

        # Loại bỏ các messages về absl::InitializeLog
        if "absl::initializelog" in msg_lower or "absl::InitializeLog" in msg:
            return False

        # Loại bỏ các messages về ALTS creds (nhiều pattern khác nhau)
        if any(
            pattern in msg
            for pattern in [
                "ALTS creds",
                "alts_credentials",
                "alts creds ignored",
                "not running on gcp",
                "untrusted alts is not enabled",
            ]
        ):
            return False

        # Loại bỏ các messages từ absl logger
        if record.name.startswith("absl.") or "absl" in record.name.lower():
            return False

        # Loại bỏ messages có pattern E0000 từ gRPC/absl
        if msg.startswith("E0000") and ("alts" in msg_lower or "cred" in msg_lower):
            return False

        return True


_stderr_filter_active = False


def _suppress_google_logs():
    """
    Suppress logging từ Google libraries (gRPC, absl, etc.)
    Bao gồm cả việc filter stderr trực tiếp.
    """
    global _stderr_filter_active

    # Set environment variables
    os.environ["GRPC_VERBOSITY"] = "ERROR"
    os.environ["GLOG_minloglevel"] = "2"
    os.environ["GRPC_PYTHON_LOG_LEVEL"] = "ERROR"
    os.environ["TF_CPP_MIN_LOG_LEVEL"] = "2"

    # Suppress absl logging nếu có
    try:
        import absl.logging

        absl.logging.set_verbosity(absl.logging.ERROR)
        # Disable absl handler
        for handler in absl.logging._absl_logger.handlers:
            handler.setLevel(logging.ERROR)
    except Exception:
        pass

    # Filter stderr và stdout để chặn messages in trực tiếp (re-apply để đảm bảo)
    global _stderr_filter_active, _stdout_filter_active
    try:
        # Kiểm tra nếu đã là NoisyMessageFilter rồi thì không cần apply lại
        if not isinstance(sys.stderr, NoisyMessageFilter):
            original_stderr = (
                sys.stderr
                if not isinstance(sys.stderr, NoisyMessageFilter)
                else sys.stderr.original_stream
            )
            sys.stderr = NoisyMessageFilter(original_stderr)
            _stderr_filter_active = True
    except Exception:
        pass

    try:
        if not isinstance(sys.stdout, NoisyMessageFilter):
            original_stdout = (
                sys.stdout
                if not isinstance(sys.stdout, NoisyMessageFilter)
                else sys.stdout.original_stream
            )
            sys.stdout = NoisyMessageFilter(original_stdout)
            _stdout_filter_active = True
    except Exception:
        pass

    # Apply filter cho root logger và các loggers cụ thể
    google_filter = GoogleLogFilter()
    root_logger = logging.getLogger()
    root_logger.addFilter(google_filter)

    # Suppress logging từ các Google libraries
    for lib_name in [
        "google",
        "grpc",
        "absl",
        "google.generativeai",
        "google.api_core",
        "google.auth",
        "grpc._cython",
    ]:
        lib_logger = logging.getLogger(lib_name)
        lib_logger.setLevel(logging.ERROR)
        lib_logger.propagate = False
        lib_logger.addFilter(google_filter)


def _parse_pages(pages_str: str) -> Optional[List[int]]:
    """
    Parse chuỗi pages thành danh sách số trang.
    Hỗ trợ:
    - "1,2,5,7" → [1, 2, 5, 7]
    - "1-7" → [1, 2, 3, 4, 5, 6, 7]
    - "1-3,5,7-9" → [1, 2, 3, 5, 7, 8, 9]

    Returns: List[int] hoặc None nếu không hợp lệ
    """
    if not pages_str or not pages_str.strip():
        return None

    pages_str = pages_str.strip()
    # Loại bỏ dấu ngoặc vuông hoặc ngoặc tròn nếu có (để tương thích ngược)
    if (pages_str.startswith("[") and pages_str.endswith("]")) or (
        pages_str.startswith("(") and pages_str.endswith(")")
    ):
        pages_str = pages_str[1:-1].strip()

    pages: List[int] = []
    parts = [p.strip() for p in pages_str.split(",")]

    for part in parts:
        part = part.strip()
        if not part:
            continue

        # Kiểm tra có phải range không (ví dụ: "1-7")
        if "-" in part:
            try:
                start, end = part.split("-", 1)
                start = int(start.strip())
                end = int(end.strip())
                if start > end:
                    logger.warning(f"Range không hợp lệ: {part} (start > end). Bỏ qua.")
                    continue
                pages.extend(range(start, end + 1))
            except ValueError:
                logger.warning(f"Range không hợp lệ: {part}. Bỏ qua.")
                continue
        else:
            # Số trang đơn lẻ
            try:
                page_num = int(part)
                if page_num > 0:
                    pages.append(page_num)
            except ValueError:
                logger.warning(f"Số trang không hợp lệ: {part}. Bỏ qua.")
                continue

    # Loại bỏ trùng lặp và sắp xếp
    pages = sorted(list(set(pages)))

    if not pages:
        logger.warning("Không có trang hợp lệ nào được parse.")
        return None

    return pages


def _ensure_logger_config() -> None:
    """Đảm bảo logger có handler để in ra console và lưu file khi chạy trực tiếp.
    Tránh tình trạng không thấy log do thiếu cấu hình bên ngoài.
    """
    if getattr(_ensure_logger_config, "_configured", False):
        return

    # Suppress Google logs trước khi cấu hình logger
    _suppress_google_logs()

    logger.setLevel(logging.INFO)
    logger.propagate = False
    # Kiểm tra có StreamHandler/FileHandler chưa
    has_stream = any(isinstance(h, logging.StreamHandler) for h in logger.handlers)
    has_file = any(isinstance(h, logging.FileHandler) for h in logger.handlers)
    # Console handler
    if not has_stream:
        ch = logging.StreamHandler()
        ch.setLevel(logging.INFO)
        ch.setFormatter(logging.Formatter("%(asctime)s - %(levelname)s - %(message)s"))
        ch.addFilter(GoogleLogFilter())  # Apply filter để loại bỏ Google logs
        logger.addHandler(ch)
    # File handler
    if not has_file:
        try:
            os.makedirs("logs", exist_ok=True)
            fh = logging.FileHandler(
                os.path.join("logs", "ocr_runtime.log"), encoding="utf-8"
            )
            fh.setLevel(logging.INFO)
            fh.setFormatter(
                logging.Formatter("%(asctime)s - %(levelname)s - %(message)s")
            )
            fh.addFilter(GoogleLogFilter())  # Apply filter để loại bỏ Google logs
            logger.addHandler(fh)
        except Exception:
            pass
    setattr(_ensure_logger_config, "_configured", True)


def _load_yaml(path: str) -> dict:
    if not os.path.exists(path):
        raise FileNotFoundError(path)
    with open(path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f) or {}


def load_ocr_config(config_path: str = "config/config.yaml") -> dict:
    cfg = _load_yaml(config_path)
    ocr_cfg = cfg.get("ocr") or {}
    # Lưu config_path để dùng sau
    ocr_cfg["_config_path"] = config_path
    # Lưu api_keys từ root config để dùng cho AI cleanup
    ocr_cfg["_root_api_keys"] = cfg.get("api_keys", [])
    # Lưu safety_level từ root config (nếu có) để dùng cho AI cleanup/spell check
    # Ưu tiên: ocr.safety_level > root safety_level > default BLOCK_ONLY_HIGH
    if "safety_level" not in ocr_cfg:
        ocr_cfg["safety_level"] = cfg.get("safety_level", "BLOCK_ONLY_HIGH")
    return ocr_cfg


def _build_safety_settings(safety_level: str = "BLOCK_ONLY_HIGH") -> List[dict]:
    """
    Tạo safety settings cho Google Gemini API.
    Học hỏi từ module dịch thuật (model_router.py).

    Args:
        safety_level: Safety level từ config (BLOCK_NONE, BLOCK_ONLY_HIGH, BLOCK_MEDIUM_AND_ABOVE, BLOCK_LOW_AND_ABOVE)

    Returns:
        List of safety settings dicts cho GenerativeModel
    """
    safety_level = safety_level.upper() if safety_level else "BLOCK_ONLY_HIGH"

    # Các levels hợp lệ từ Google Gemini API
    valid_levels = [
        "BLOCK_NONE",
        "BLOCK_ONLY_HIGH",
        "BLOCK_MEDIUM_AND_ABOVE",
        "BLOCK_LOW_AND_ABOVE",
    ]
    if safety_level not in valid_levels:
        logger.warning(
            f"Safety level '{safety_level}' không hợp lệ. Dùng default: BLOCK_ONLY_HIGH"
        )
        safety_level = "BLOCK_ONLY_HIGH"

    return [
        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": safety_level},
        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": safety_level},
        {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": safety_level},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": safety_level},
    ]


def _bundle_base_dir() -> str:
    """Return base dir for bundled resources (PyInstaller) or script dir."""
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        return getattr(sys, "_MEIPASS")  # PyInstaller temp dir
    # Fallback: repo/script directory
    return os.path.dirname(os.path.abspath(sys.argv[0]))


def _detect_bundled_binaries(ocr_cfg: dict) -> dict:
    """
    If config values are missing, try to detect bundled Tesseract and Poppler paths.
    - Looks for vendor/tesseract/tesseract.exe
    - Looks for vendor/poppler/bin
    """
    cfg = dict(ocr_cfg) if ocr_cfg else {}
    base = _bundle_base_dir()
    # Detect Tesseract
    if not cfg.get("tesseract_cmd"):
        cand = os.path.join(base, "tesseract", "tesseract.exe")
        if not os.path.exists(cand):
            cand = os.path.join(base, "vendor", "tesseract", "tesseract.exe")
        if os.path.exists(cand):
            cfg["tesseract_cmd"] = cand.replace("\\", "/")
    # Detect Poppler bin
    if not cfg.get("poppler_path"):
        cand_dir = os.path.join(base, "poppler", "bin")
        if not os.path.isdir(cand_dir):
            cand_dir = os.path.join(base, "vendor", "poppler", "bin")
        if os.path.isdir(cand_dir):
            cfg["poppler_path"] = cand_dir.replace("\\", "/")
    return cfg


def _pip_install(package: str) -> None:
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])
    except Exception as e:
        raise RuntimeError(f"Không thể cài gói '{package}': {e}")


def _ensure_dependencies(ocr_cfg: dict) -> None:
    global \
        Image, \
        pytesseract, \
        convert_from_path, \
        tqdm, \
        PyPDF2, \
        pdfplumber, \
        fitz, \
        Document, \
        Inches, \
        Pt, \
        WD_PARAGRAPH_ALIGNMENT
    # Pillow
    if Image is None:
        _pip_install("Pillow>=10.0.0")
        from PIL import Image as _Image

        Image = _Image
    # pdf2image
    if convert_from_path is None:
        _pip_install("pdf2image>=1.17.0")
        from pdf2image import convert_from_path as _convert

        convert_from_path = _convert
    # pytesseract (runtime still needs system Tesseract)
    if pytesseract is None:
        _pip_install("pytesseract>=0.3.10")
        import pytesseract as _p

        pytesseract = _p
    # pdfplumber (preferred for text extraction)
    if pdfplumber is None:
        try:
            _pip_install("pdfplumber>=0.9.0")
            import pdfplumber as _pp

            pdfplumber = _pp
        except Exception:
            pass
    # PyPDF2 (fallback for text extraction)
    if PyPDF2 is None:
        try:
            _pip_install("PyPDF2>=3.0.0")
            import PyPDF2 as _pypdf

            PyPDF2 = _pypdf
        except Exception:
            pass
    # yaml (PyYAML) was already imported to read config; skip
    # tqdm optional
    if bool(ocr_cfg.get("show_progress", True)) and tqdm is None:
        try:
            _pip_install("tqdm>=4.65.0")
            from tqdm import tqdm as _tqdm

            tqdm = _tqdm
        except Exception:
            tqdm = None
    # PyMuPDF (for extracting images from PDF)
    if fitz is None:
        try:
            _pip_install("PyMuPDF>=1.23.0")
            import fitz as _fitz

            fitz = _fitz
        except Exception:
            fitz = None
    # python-docx (for creating DOCX output)
    if Document is None:
        try:
            _pip_install("python-docx>=1.0.0")
            from docx import Document as _Document
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as _WD_PARAGRAPH_ALIGNMENT
            from docx.shared import Inches as _Inches
            from docx.shared import Pt as _Pt

            Document = _Document
            global Inches, Pt, WD_PARAGRAPH_ALIGNMENT
            Inches = _Inches
            Pt = _Pt
            WD_PARAGRAPH_ALIGNMENT = _WD_PARAGRAPH_ALIGNMENT
        except Exception:
            Document = None
            Inches = None
            Pt = None
            WD_PARAGRAPH_ALIGNMENT = None


def _apply_tesseract_cfg(ocr_cfg: dict) -> None:
    if pytesseract is None:
        raise RuntimeError(
            "pytesseract not installed. Please install pytesseract and system Tesseract."
        )
    # Allow auto-detect of bundled binaries
    _cfg = _detect_bundled_binaries(ocr_cfg)
    tesseract_cmd = _cfg.get("tesseract_cmd")
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd


def _normalize_lang_code(lang: str) -> str:
    """
    Chuyển đổi mã ngôn ngữ từ format ngắn (VN, EN, CN) sang Tesseract format (vie, eng, chi).
    Hỗ trợ backward compatibility với format cũ.

    Args:
        lang: Language string có thể là "VN", "EN", "CN", "auto", hoặc format cũ "vie", "eng", "chi"

    Returns:
        Tesseract language code hoặc "auto"
    """
    if not lang:
        return "vie"

    lang = lang.strip().upper()

    # Mapping từ format ngắn sang Tesseract
    lang_map = {
        "VN": "vie",
        "EN": "eng",
        "CN": "chi",
        "VIE": "vie",  # Backward compatibility
        "ENG": "eng",  # Backward compatibility
        "CHI": "chi",  # Backward compatibility
        "AUTO": "auto",
    }

    # Xử lý kết hợp ngôn ngữ (VD: "VN+EN" hoặc "vie+eng")
    if "+" in lang:
        parts = lang.split("+")
        normalized_parts = []
        for part in parts:
            part = part.strip().upper()
            normalized = lang_map.get(
                part, part.lower()
            )  # Fallback về lowercase nếu không map được
            normalized_parts.append(normalized)
        return "+".join(normalized_parts)

    # Xử lý single language
    return lang_map.get(lang, lang.lower())  # Fallback về lowercase nếu không map được


def _is_cjk_character(char: str) -> bool:
    """
    Kiểm tra xem ký tự có phải là CJK (Chinese, Japanese, Korean) không.
    Dựa trên Unicode ranges cho CJK.
    """
    if not char:
        return False
    code = ord(char)
    # CJK Unified Ideographs: U+4E00–U+9FFF
    # CJK Extension A: U+3400–U+4DBF
    # CJK Extension B: U+20000–U+2A6DF
    # CJK Compatibility: U+F900–U+FAFF
    return (
        0x4E00 <= code <= 0x9FFF  # CJK Unified Ideographs
        or 0x3400 <= code <= 0x4DBF  # CJK Extension A
        or 0xF900 <= code <= 0xFAFF  # CJK Compatibility
    )


def _count_cjk_characters(text: str) -> int:
    """Đếm số ký tự CJK trong text."""
    return sum(1 for char in text if _is_cjk_character(char))


# DEPRECATED: Auto-detect ngôn ngữ đã bị loại bỏ do kém hiệu quả
# Hàm này không còn được sử dụng trong code chính
def _detect_language_from_image(
    img: "Image.Image", ocr_cfg: dict, candidate_langs: List[str] = None
) -> tuple[str, float]:
    """
    [DEPRECATED] Tự động phát hiện ngôn ngữ từ ảnh bằng cách OCR với nhiều ngôn ngữ và so sánh confidence.
    Giới hạn với 3 ngôn ngữ: VN (vie), EN (eng), CN (chi_sim + chi_tra).

    Cải thiện: Pre-check CJK characters để ưu tiên Chinese detection.

    LƯU Ý: Hàm này đã bị loại bỏ do kém hiệu quả trong thực tế.

    Args:
        img: PIL Image object
        ocr_cfg: OCR config
        candidate_langs: Danh sách ngôn ngữ candidate để test
                        (mặc định: ["vie", "eng", "chi_sim", "chi_tra"] - chỉ VN, EN, CN)

    Returns:
        Tuple (Tesseract language code, score) được detect (e.g., ("chi_sim", 85.5))
        Nếu không đủ tin cậy (score < 30), trả về ("vie", score) làm fallback
    """
    if candidate_langs is None:
        # Giới hạn chỉ với VN, EN, CN (bao gồm cả chi_sim và chi_tra)
        candidate_langs = ["vie", "eng", "chi_sim", "chi_tra"]

    psm = int(ocr_cfg.get("psm", 3) or 3)
    config = f"--psm {psm}"

    # PRE-CHECK: OCR với Chinese languages để detect CJK characters
    # Nếu có nhiều CJK characters, ưu tiên Chinese detection và loại bỏ non-CJK languages
    cjk_char_count = 0
    try:
        # Test với cả chi_sim và chi_tra để chắc chắn
        for pre_lang in ["chi_sim", "chi_tra"]:
            try:
                quick_data = pytesseract.image_to_data(
                    img,
                    lang=pre_lang,
                    config=config,
                    output_type=pytesseract.Output.DICT,
                )
                quick_text = "".join(
                    [
                        quick_data["text"][i]
                        for i in range(len(quick_data["text"]))
                        if quick_data["text"][i].strip()
                        and int(quick_data["conf"][i]) > 25
                    ]
                )  # Giảm threshold xuống 25
                quick_cjk_count = _count_cjk_characters(quick_text)
                if quick_cjk_count > cjk_char_count:
                    cjk_char_count = quick_cjk_count
            except Exception:
                continue

        logger.debug(
            f"Pre-check: Phát hiện {cjk_char_count} ký tự CJK trong text (từ pre-check Chinese)"
        )
    except Exception:
        pass

    best_lang = "vie"  # Default fallback
    best_score = 0
    min_confidence_threshold = 30.0  # Score tối thiểu để được coi là đáng tin cậy

    # Nếu có CJK characters (>5), ưu tiên Chinese languages (giảm threshold để nhạy hơn)
    has_cjk = cjk_char_count > 5
    if has_cjk:
        logger.info(
            f"Phát hiện {cjk_char_count} ký tự CJK, ưu tiên Chinese detection (bỏ qua VN/EN)"
        )
        # Chỉ test với Chinese languages nếu có CJK
        candidate_langs = [lang for lang in candidate_langs if "chi" in lang.lower()]
        if not candidate_langs:
            candidate_langs = ["chi_sim", "chi_tra"]  # Fallback

    for lang in candidate_langs:
        try:
            data = pytesseract.image_to_data(
                img, lang=lang, config=config, output_type=pytesseract.Output.DICT
            )
            confidences = [int(conf) for conf in data["conf"] if int(conf) > 0]

            if not confidences:
                continue

            avg_conf = sum(confidences) / len(confidences)

            # Tính số ký tự thực sự được nhận dạng và số ký tự CJK
            recognized_chars = []
            recognized_text = ""
            for i in range(len(data["text"])):
                text_item = data["text"][i].strip()
                conf = int(data["conf"][i])
                # Chỉ tính các ký tự có confidence > 0 và không phải chỉ là whitespace
                if text_item and conf > 0:
                    # Loại bỏ các ký tự đặc biệt không có ý nghĩa (chỉ giữ chữ, số, dấu câu cơ bản)
                    if any(c.isalnum() for c in text_item) or len(text_item) > 1:
                        recognized_chars.extend(list(text_item))
                        recognized_text += text_item

            char_count = len(recognized_chars)
            cjk_count_in_result = _count_cjk_characters(recognized_text)
            cjk_ratio = cjk_count_in_result / max(
                char_count, 1
            )  # Tỷ lệ CJK trong text nhận dạng được

            # Với CJK (Chinese), ưu tiên số ký tự cao hơn (vì mỗi ký tự CJK = 1 word)
            is_cjk = "chi" in lang.lower()

            # CRITICAL: Nếu text chứa CJK nhưng đang test non-CJK language → PENALTY lớn (giảm threshold)
            if not is_cjk and cjk_ratio > 0.2:  # Nếu >20% là CJK (giảm từ 30%)
                logger.debug(
                    f"  {lang}: SKIP - Text chứa {cjk_ratio * 100:.1f}% CJK nhưng đang test non-CJK language"
                )
                continue  # Bỏ qua hoàn toàn nếu có CJK

            # CRITICAL: Nếu pre-check phát hiện CJK và đang test non-CJK language → SKIP
            if not is_cjk and cjk_char_count > 3:  # Giảm threshold xuống 3 (từ 10)
                logger.debug(
                    f"  {lang}: SKIP - Pre-check phát hiện {cjk_char_count} CJK, bỏ qua non-CJK language"
                )
                continue  # Bỏ qua hoàn toàn

            # Scoring cải thiện với normalization tốt hơn và ưu tiên CJK:
            # - Với CJK: character count quan trọng hơn (40% conf, 60% chars) - tăng mạnh trọng số chars
            # - Với non-CJK: confidence quan trọng hơn (80% conf, 20% chars)
            # - Normalize char_count: với CJK max ~400 chars/page (thực tế), non-CJK max ~800 chars/page
            if is_cjk:
                char_max = 400.0
                char_weight = 0.60  # Tăng mạnh trọng số cho chars (từ 0.55)
                conf_weight = 0.40

                # MAJOR BONUS cho CJK nếu có nhiều CJK characters được nhận dạng
                if cjk_count_in_result > 20:  # Nếu có >20 CJK chars
                    cjk_bonus = min(
                        (cjk_count_in_result / 100.0) * 25, 25.0
                    )  # Bonus lớn: tối đa 25 điểm
                    logger.debug(
                        f"  {lang}: CJK bonus = {cjk_bonus:.1f} (có {cjk_count_in_result} CJK chars)"
                    )
                else:
                    cjk_bonus = 0

                # Bonus nếu có nhiều ký tự tổng thể
                if char_count > (char_max * 0.2):  # >20% ký tự (giảm threshold)
                    char_bonus = min(
                        (char_count / char_max) * 20, 20.0
                    )  # Bonus tối đa 20 điểm (tăng từ 15)
                else:
                    char_bonus = 0
            else:
                char_max = 800.0
                char_weight = 0.20  # Giảm trọng số chars (từ 0.25)
                conf_weight = 0.80  # Tăng trọng số confidence (từ 0.75)
                cjk_bonus = 0

                # Bonus nhỏ hơn cho non-CJK
                if char_count > (char_max * 0.3):  # >30% ký tự
                    char_bonus = min(
                        (char_count / char_max) * 8, 8.0
                    )  # Bonus tối đa 8 điểm (giảm từ 10)
                else:
                    char_bonus = 0

            char_score = min(char_count / char_max, 1.0) * 100
            base_score = avg_conf * conf_weight + char_score * char_weight
            score = base_score + char_bonus + cjk_bonus

            # MAJOR PENALTY: Nếu confidence quá thấp (< 30), giảm score đáng kể
            if avg_conf < 30:
                score *= 0.4  # Giảm mạnh hơn: 60% (từ 50%)

            # MAJOR PENALTY: Nếu char_count quá ít (< 10 ký tự), có thể là noise
            if char_count < 10:
                score *= 0.5  # Giảm mạnh hơn: 50% (từ 40%)

            # CRITICAL: Nếu có CJK trong result nhưng đang test non-CJK (should not happen sau filter, nhưng double-check)
            if not is_cjk and cjk_count_in_result > 5:
                score *= 0.3  # Giảm mạnh: chỉ còn 30% score

            # BONUS: Nếu có CJK và đang test CJK language → bonus thêm
            if is_cjk and cjk_count_in_result > 0:
                score += min(
                    cjk_count_in_result * 0.5, 10.0
                )  # Bonus thêm: 0.5 điểm mỗi CJK char, tối đa 10

            logger.debug(
                f"  {lang}: conf={avg_conf:.1f}, chars={char_count} (CJK: {cjk_count_in_result}), "
                f"base={base_score:.1f}, char_bonus={char_bonus:.1f}, cjk_bonus={cjk_bonus:.1f}, final={score:.1f}"
            )

            if score > best_score:
                best_score = score
                best_lang = lang
        except Exception as e:
            logger.debug(f"  {lang}: Failed ({e})")
            continue

    # CRITICAL VALIDATION: Nếu detect ra non-CJK nhưng có nhiều CJK trong pre-check → Không tin cậy
    is_detected_cjk = "chi" in best_lang.lower()
    if not is_detected_cjk and cjk_char_count > 10:
        logger.warning(
            f"VALIDATION FAIL: Detect ra '{best_lang}' (non-CJK) nhưng pre-check phát hiện {cjk_char_count} CJK chars. "
            f"Có thể nhầm lẫn. Thử lại với chỉ Chinese languages..."
        )
        # Force retry chỉ với Chinese languages
        try:
            chinese_only_results = []
            for lang in ["chi_sim", "chi_tra"]:
                try:
                    data = pytesseract.image_to_data(
                        img,
                        lang=lang,
                        config=config,
                        output_type=pytesseract.Output.DICT,
                    )
                    confidences = [int(conf) for conf in data["conf"] if int(conf) > 0]
                    if not confidences:
                        continue
                    avg_conf = sum(confidences) / len(confidences)
                    recognized_text = "".join(
                        [
                            data["text"][i]
                            for i in range(len(data["text"]))
                            if data["text"][i].strip() and int(data["conf"][i]) > 0
                        ]
                    )
                    cjk_count = _count_cjk_characters(recognized_text)
                    char_count = len(recognized_text)

                    # Scoring cho Chinese
                    char_score = min(char_count / 400.0, 1.0) * 100
                    score = avg_conf * 0.4 + char_score * 0.6
                    if cjk_count > 20:
                        score += min((cjk_count / 100.0) * 25, 25.0)
                    if cjk_count > 0:
                        score += min(cjk_count * 0.5, 10.0)

                    chinese_only_results.append((lang, score, cjk_count))
                    logger.debug(
                        f"  Retry {lang}: conf={avg_conf:.1f}, chars={char_count} (CJK: {cjk_count}), score={score:.1f}"
                    )
                except Exception:
                    continue

            if chinese_only_results:
                # Chọn Chinese language có score cao nhất
                best_chinese = max(chinese_only_results, key=lambda x: x[1])
                logger.info(
                    f"VALIDATION PASS: Sau retry, chọn '{best_chinese[0]}' với score {best_chinese[1]:.1f} "
                    f"(có {best_chinese[2]} CJK chars)"
                )
                return (best_chinese[0], best_chinese[1])
            else:
                logger.warning(
                    f"Retry Chinese languages không thành công. Giữ nguyên kết quả '{best_lang}' nhưng không tin cậy."
                )
        except Exception as e:
            logger.warning(f"Validation retry failed: {e}")

    # Validation: Kiểm tra xem score có đủ tin cậy không
    if best_score < min_confidence_threshold:
        logger.warning(
            f"Auto-detect score quá thấp ({best_score:.1f} < {min_confidence_threshold}), có thể không chính xác."
        )
        # Nếu có CJK nhưng detect non-CJK, không tin cậy
        if not is_detected_cjk and cjk_char_count > 5:
            logger.warning(
                f"Cảnh báo: Có {cjk_char_count} CJK chars nhưng detect ra '{best_lang}'. Kết quả không tin cậy."
            )

    logger.debug(f"Auto-detect: Chọn {best_lang} với score {best_score:.1f}")
    return (best_lang, best_score)


# DEPRECATED: Auto-detect ngôn ngữ đã bị loại bỏ do kém hiệu quả
# Hàm này không còn được sử dụng trong code chính
def _detect_language_from_multiple_pages(
    page_paths: List[Path], ocr_cfg: dict, max_sample_pages: int = 5
) -> str:
    """
    [DEPRECATED] Detect ngôn ngữ từ nhiều trang để tránh nhầm lẫn với trang đầu ít chữ.
    Sample nhiều trang và chọn kết quả có confidence cao nhất.
    Giới hạn chỉ với 3 ngôn ngữ: VN (vie), EN (eng), CN (chi_sim + chi_tra).

    LƯU Ý: Hàm này đã bị loại bỏ do kém hiệu quả trong thực tế.

    Args:
        page_paths: Danh sách đường dẫn đến các trang (ảnh)
        ocr_cfg: OCR config
        max_sample_pages: Số trang tối đa để sample (mặc định: 5)

    Returns:
        Tesseract language code được detect (vie, eng, chi_sim, hoặc chi_tra)
        Fallback về "vie" nếu không đủ tin cậy
    """
    # Giới hạn candidate languages chỉ với VN, EN, CN
    candidate_langs_auto = ["vie", "eng", "chi_sim", "chi_tra"]
    if not page_paths:
        logger.warning("Không có trang nào để sample, dùng mặc định 'vie'")
        return "vie"

    # Sample nhiều trang (tránh chỉ lấy trang đầu)
    num_pages = len(page_paths)
    sample_size = min(max_sample_pages, num_pages)

    # Chiến lược sampling (ưu tiên tránh trang đầu có thể là bìa):
    # - Nếu <= 3 trang: lấy tất cả
    # - Nếu 4-5 trang: lấy từ trang 2 (bỏ qua trang 1 có thể là bìa)
    # - Nếu > 5 trang: lấy 2-3 trang sau trang đầu + 1-2 trang giữa
    if num_pages <= 3:
        pages_to_sample = list(range(num_pages))
    elif num_pages <= 5:
        # Bỏ qua trang đầu (index 0), lấy từ trang 2 trở đi
        pages_to_sample = list(range(1, min(sample_size + 1, num_pages)))
    else:
        # Với PDF lớn: bỏ qua trang đầu, lấy 2-3 trang tiếp theo + 1-2 trang giữa
        pages_to_sample = [1, 2]  # Trang 2, 3 (bỏ qua trang 1)
        if num_pages > 7:
            # Thêm 2 trang ở giữa
            mid_start = num_pages // 3
            pages_to_sample.extend([mid_start, mid_start + 1])
        else:
            # Thêm 1 trang giữa
            pages_to_sample.append(num_pages // 2)
        pages_to_sample = pages_to_sample[:sample_size]

    # Map indices để hiển thị số trang (1-based)
    page_numbers = [p + 1 for p in pages_to_sample]
    logger.info(
        f"Auto-detect language từ {len(pages_to_sample)} trang (số trang: {page_numbers})..."
    )

    results: List[tuple[str, float]] = []  # [(lang, score), ...]

    for idx_in_list in pages_to_sample:
        if idx_in_list >= len(page_paths):
            continue
        page_path = page_paths[idx_in_list]
        try:
            with Image.open(str(page_path)) as img:
                # Chỉ test với 3 ngôn ngữ: VN, EN, CN
                lang, score = _detect_language_from_image(
                    img, ocr_cfg, candidate_langs=candidate_langs_auto
                )
                results.append((lang, score))
                logger.debug(f"  Trang {idx_in_list + 1}: {lang} (score: {score:.1f})")
        except Exception as e:
            logger.debug(f"  Trang {idx_in_list + 1}: Failed ({e})")
            continue

    if not results:
        logger.warning("Không thể detect từ bất kỳ trang nào, dùng mặc định 'vie'")
        return "vie"

    # Chọn language có score cao nhất với validation chặt chẽ hơn
    lang_scores: dict[str, List[float]] = {}
    for lang, score in results:
        if lang not in lang_scores:
            lang_scores[lang] = []
        lang_scores[lang].append(score)

    # Tính average score và consistency cho mỗi language
    best_lang = "vie"
    best_final_score = 0
    min_avg_score_threshold = 25.0  # Threshold tối thiểu cho average score
    min_consistency_ratio = 0.5  # Ít nhất 50% trang phải detect cùng language

    for lang, scores in lang_scores.items():
        avg_score = sum(scores) / len(scores)
        consistency_ratio = len(scores) / len(pages_to_sample)

        # Chỉ xét nếu consistency đủ cao và avg_score đủ cao
        if consistency_ratio < min_consistency_ratio:
            logger.debug(
                f"  {lang}: Bỏ qua (consistency {consistency_ratio:.2f} < {min_consistency_ratio})"
            )
            continue

        if avg_score < min_avg_score_threshold:
            logger.debug(
                f"  {lang}: Bỏ qua (avg_score {avg_score:.1f} < {min_avg_score_threshold})"
            )
            continue

        # Consistency bonus: xuất hiện càng nhiều càng tốt (tối đa 15 điểm)
        consistency_bonus = consistency_ratio * 15

        # Quality bonus: avg_score càng cao càng tốt (tối đa 10 điểm)
        quality_bonus = min((avg_score / 100.0) * 10, 10.0)

        final_score = avg_score + consistency_bonus + quality_bonus

        logger.debug(
            f"  {lang}: avg={avg_score:.1f}, consistency={consistency_ratio:.2f} ({len(scores)}/{len(pages_to_sample)}), "
            f"consistency_bonus={consistency_bonus:.1f}, quality_bonus={quality_bonus:.1f}, final={final_score:.1f}"
        )

        if final_score > best_final_score:
            best_final_score = final_score
            best_lang = lang

    # Validation cuối cùng: Nếu không có language nào đạt threshold, dùng fallback
    if best_final_score == 0:
        logger.warning(
            f"Không có language nào đạt threshold (min_avg={min_avg_score_threshold}, min_consistency={min_consistency_ratio}). Dùng fallback 'vie'."
        )
        return "vie"

    logger.info(
        f"Auto-detected language: {best_lang} (final_score: {best_final_score:.1f}, từ {len(results)}/{len(pages_to_sample)} trang thành công)"
    )
    return best_lang


def _detect_chinese_variant(img: "Image.Image", ocr_cfg: dict) -> str:
    """
    Tự động nhận biết tiếng Trung giản thể hay phồn thể.
    Returns: "chi_sim" hoặc "chi_tra"
    """
    psm = int(ocr_cfg.get("psm", 3) or 3)
    config = f"--psm {psm}"

    try:
        # OCR với cả 2 ngôn ngữ và so sánh confidence
        # Simplified Chinese
        data_sim = pytesseract.image_to_data(
            img, lang="chi_sim", config=config, output_type=pytesseract.Output.DICT
        )
        confidences_sim = [int(conf) for conf in data_sim["conf"] if int(conf) > 0]
        avg_conf_sim = (
            sum(confidences_sim) / len(confidences_sim) if confidences_sim else 0
        )
        # Đếm số ký tự được nhận dạng (có confidence > 0)
        char_count_sim = sum(
            1
            for i, text_item in enumerate(data_sim["text"])
            if text_item.strip() and int(data_sim["conf"][i]) > 0
        )

        # Traditional Chinese
        data_tra = pytesseract.image_to_data(
            img, lang="chi_tra", config=config, output_type=pytesseract.Output.DICT
        )
        confidences_tra = [int(conf) for conf in data_tra["conf"] if int(conf) > 0]
        avg_conf_tra = (
            sum(confidences_tra) / len(confidences_tra) if confidences_tra else 0
        )
        # Đếm số ký tự được nhận dạng (có confidence > 0)
        char_count_tra = sum(
            1
            for i, text_item in enumerate(data_tra["text"])
            if text_item.strip() and int(data_tra["conf"][i]) > 0
        )

        # Quyết định dựa trên confidence và số ký tự
        # Ưu tiên confidence, nếu gần bằng nhau thì ưu tiên số ký tự nhiều hơn
        score_sim = (
            avg_conf_sim * 0.7
            + (char_count_sim / max(char_count_sim + char_count_tra, 1)) * 30 * 0.3
        )
        score_tra = (
            avg_conf_tra * 0.7
            + (char_count_tra / max(char_count_sim + char_count_tra, 1)) * 30 * 0.3
        )

        if score_sim > score_tra:
            detected = "chi_sim"
            logger.debug(
                f"Chinese variant detected: Simplified (conf: {avg_conf_sim:.1f}, chars: {char_count_sim})"
            )
        else:
            detected = "chi_tra"
            logger.debug(
                f"Chinese variant detected: Traditional (conf: {avg_conf_tra:.1f}, chars: {char_count_tra})"
            )

        return detected
    except Exception as e:
        # Fallback: mặc định là Simplified (phổ biến hơn)
        logger.warning(f"Không thể detect Chinese variant: {e}. Mặc định dùng chi_sim")
        return "chi_sim"


def _resolve_language(
    lang: str, ocr_cfg: dict, sample_img: Optional["Image.Image"] = None
) -> str:
    """
    Resolve language code, chỉ hỗ trợ Chinese variant detection (giản thể/phồn thể).
    Auto-detect ngôn ngữ đã được loại bỏ do kém hiệu quả.

    Args:
        lang: Language string từ config (có thể là "VN", "EN", "CN", "VN+EN", "chi", "chi_sim", "chi_tra", etc.)
        ocr_cfg: OCR config
        sample_img: Optional sample image để detect Chinese variant (chỉ khi lang="CN" hoặc "chi")

    Returns:
        Resolved language string cho Tesseract (e.g., "chi_sim", "chi_tra", "vie+eng")
    """
    if not lang:
        return "vie"

    # Normalize: VN/EN/CN → vie/eng/chi
    lang = _normalize_lang_code(lang)

    # Loại bỏ auto-detect: nếu config là "auto", cảnh báo và fallback về "vie"
    if lang == "auto" or lang.startswith("auto+"):
        logger.warning(
            f"Auto-detect ngôn ngữ đã bị loại bỏ do kém hiệu quả. "
            f"Config '{lang}' không được hỗ trợ. Vui lòng chỉ định rõ ngôn ngữ (VN/EN/CN). "
            f"Fallback về 'vie'."
        )
        lang = "vie"

    # Chỉ hỗ trợ detect Chinese variant (giản thể/phồn thể) khi lang="CN" hoặc "chi"
    # Kiểm tra nếu có "chi" (cần detect variant: Simplified vs Traditional)
    if "chi" in lang.lower() and "chi_sim" not in lang and "chi_tra" not in lang:
        # Cần detect variant
        if sample_img is not None:
            detected_variant = _detect_chinese_variant(sample_img, ocr_cfg)
            # Replace "chi" bằng variant detected
            lang = lang.replace("chi", detected_variant).replace(
                "Chi", detected_variant
            )
            # Clean up duplicate "+" nếu có
            lang = lang.replace(
                f"{detected_variant}+{detected_variant}", detected_variant
            )
            logger.info(
                f"Auto-detected Chinese variant: {detected_variant} → Language: {lang}"
            )
        else:
            # Không có sample image → mặc định Simplified
            detected_variant = "chi_sim"
            lang = lang.replace("chi", detected_variant).replace(
                "Chi", detected_variant
            )
            logger.info(
                f"No sample image for detection, defaulting to chi_sim → Language: {lang}"
            )

    return lang


def _image_to_text(
    img: "Image.Image", ocr_cfg: dict, lang_override: Optional[str] = None
) -> str:
    """
    OCR một ảnh thành text.

    Args:
        img: PIL Image object
        ocr_cfg: OCR config dictionary
        lang_override: Optional resolved language string (đã detect variant nếu cần)
    """
    lang = lang_override
    if lang is None:
        raw_lang = ocr_cfg.get("lang", "vie+eng")
        lang = _resolve_language(raw_lang, ocr_cfg, sample_img=img)

    psm = int(ocr_cfg.get("psm", 3) or 3)
    config = f"--psm {psm}"
    return pytesseract.image_to_string(img, lang=lang, config=config)


def detect_pdf_type(pdf_path: str) -> str:
    """
    Phát hiện PDF là scan hay text-based.
    Returns: "text" hoặc "scan"
    """
    if pdfplumber is not None:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                total_chars = 0
                sample_pages = min(3, len(pdf.pages))
                for i in range(sample_pages):
                    page = pdf.pages[i]
                    text = page.extract_text()
                    if text:
                        total_chars += len(text.strip())

                # Nếu có đủ text (> 100 chars trong 3 trang đầu) → text-based
                if total_chars > 100:
                    return "text"
                return "scan"
        except Exception:
            pass

    # Fallback: dùng PyPDF2
    if PyPDF2 is not None:
        try:
            with open(pdf_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                total_chars = 0
                sample_pages = min(3, len(reader.pages))
                for i in range(sample_pages):
                    page = reader.pages[i]
                    text = page.extract_text()
                    if text:
                        total_chars += len(text.strip())

                if total_chars > 100:
                    return "text"
                return "scan"
        except Exception:
            pass

    # Nếu không thể detect → giả định là scan
    logger.warning(f"Không thể detect PDF type, giả định là scan: {pdf_path}")
    return "scan"


def extract_text_from_pdf(
    pdf_path: str, ocr_cfg: dict, pages: Optional[List[int]] = None
) -> str:
    """
    Extract text từ PDF có text layer (không cần OCR).

    Args:
        pdf_path: Đường dẫn file PDF
        ocr_cfg: Config dictionary
        pages: Danh sách số trang cần extract (1-indexed). None = tất cả trang.
    """
    texts: List[str] = []

    # Ưu tiên pdfplumber (chính xác hơn)
    if pdfplumber is not None:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                total = len(pdf.pages)
                logger.info(f"Extract text: Tổng số trang: {total}")

                # Filter pages nếu có chỉ định
                if pages:
                    # Validate pages (phải trong khoảng [1, total])
                    valid_pages = [p for p in pages if 1 <= p <= total]
                    invalid_pages = [p for p in pages if p < 1 or p > total]
                    if invalid_pages:
                        logger.warning(
                            f"Các trang không hợp lệ (nằm ngoài 1-{total}): {invalid_pages}. Bỏ qua."
                        )
                    if not valid_pages:
                        logger.error("Không có trang hợp lệ nào để extract.")
                        return ""
                    logger.info(
                        f"Extract text: Chỉ extract {len(valid_pages)} trang: {valid_pages}"
                    )
                    pages_to_extract = sorted(set(valid_pages))
                else:
                    pages_to_extract = list(range(1, total + 1))

                show_progress = bool(ocr_cfg.get("show_progress", True))

                if show_progress and tqdm is not None and len(pages_to_extract) > 1:
                    for page_num in tqdm(
                        pages_to_extract, desc="Extract text", unit="trang"
                    ):
                        page = pdf.pages[page_num - 1]  # pdfplumber dùng 0-indexed
                        text = page.extract_text()
                        if text:
                            texts.append(text.strip())
                else:
                    for page_num in pages_to_extract:
                        page = pdf.pages[page_num - 1]  # pdfplumber dùng 0-indexed
                        text = page.extract_text()
                        if text:
                            texts.append(text.strip())
                        if len(texts) % 50 == 0:
                            logger.info(
                                f"Extract text: {len(texts)}/{len(pages_to_extract)} trang"
                            )
                return "\n\n".join(texts)
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}, trying PyPDF2...")

    # Fallback: PyPDF2
    if PyPDF2 is not None:
        try:
            with open(pdf_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                total = len(reader.pages)
                logger.info(f"Extract text: Tổng số trang: {total}")

                # Filter pages nếu có chỉ định
                if pages:
                    valid_pages = [p for p in pages if 1 <= p <= total]
                    invalid_pages = [p for p in pages if p < 1 or p > total]
                    if invalid_pages:
                        logger.warning(
                            f"Các trang không hợp lệ (nằm ngoài 1-{total}): {invalid_pages}. Bỏ qua."
                        )
                    if not valid_pages:
                        logger.error("Không có trang hợp lệ nào để extract.")
                        return ""
                    logger.info(
                        f"Extract text: Chỉ extract {len(valid_pages)} trang: {valid_pages}"
                    )
                    pages_to_extract = sorted(set(valid_pages))
                else:
                    pages_to_extract = list(range(1, total + 1))

                show_progress = bool(ocr_cfg.get("show_progress", True))

                if show_progress and tqdm is not None and len(pages_to_extract) > 1:
                    for page_num in tqdm(
                        pages_to_extract, desc="Extract text", unit="trang"
                    ):
                        page = reader.pages[page_num - 1]  # PyPDF2 dùng 0-indexed
                        text = page.extract_text()
                        if text:
                            texts.append(text.strip())
                else:
                    for page_num in pages_to_extract:
                        page = reader.pages[page_num - 1]  # PyPDF2 dùng 0-indexed
                        text = page.extract_text()
                        if text:
                            texts.append(text.strip())
                        if len(texts) % 50 == 0:
                            logger.info(
                                f"Extract text: {len(texts)}/{len(pages_to_extract)} trang"
                            )
                return "\n\n".join(texts)
        except Exception as e:
            logger.error(f"PyPDF2 failed: {e}")
            raise

    raise RuntimeError(
        "Không có thư viện extract PDF text. Cài pdfplumber hoặc PyPDF2."
    )


def extract_text_and_images_from_pdf(
    pdf_path: str, ocr_cfg: dict, pages: Optional[List[int]] = None
) -> tuple[List[dict], int]:
    """
    Extract text và images từ PDF có text layer.

    Args:
        pdf_path: Đường dẫn file PDF
        ocr_cfg: Config dictionary
        pages: Danh sách số trang cần extract (1-indexed). None = tất cả trang.

    Returns:
        tuple: (pages_data, total_pages) trong đó pages_data là list of dict với keys:
            - page_num: số trang (1-indexed)
            - text: text content
            - images: list of image data (bytes)
    """
    if fitz is None:
        raise RuntimeError(
            "PyMuPDF (fitz) chưa được cài đặt. Cài PyMuPDF để hỗ trợ extract images."
        )

    pages_data: List[dict] = []

    try:
        doc = fitz.open(pdf_path)
        total = len(doc)
        logger.info(f"Extract text và images: Tổng số trang: {total}")

        # Filter pages nếu có chỉ định
        if pages:
            valid_pages = [p for p in pages if 1 <= p <= total]
            invalid_pages = [p for p in pages if p < 1 or p > total]
            if invalid_pages:
                logger.warning(
                    f"Các trang không hợp lệ (nằm ngoài 1-{total}): {invalid_pages}. Bỏ qua."
                )
            if not valid_pages:
                logger.error("Không có trang hợp lệ nào để extract.")
                return [], 0
            logger.info(
                f"Extract text và images: Chỉ extract {len(valid_pages)} trang: {valid_pages}"
            )
            pages_to_extract = sorted(set(valid_pages))
        else:
            pages_to_extract = list(range(1, total + 1))

        show_progress = bool(ocr_cfg.get("show_progress", True))

        if show_progress and tqdm is not None and len(pages_to_extract) > 1:
            iterator = tqdm(
                pages_to_extract, desc="Extract text & images", unit="trang"
            )
        else:
            iterator = pages_to_extract

        for page_num in iterator:
            page = doc[page_num - 1]  # fitz dùng 0-indexed

            # Extract text
            text = page.get_text().strip()

            # Extract images
            images = []
            image_list = page.get_images(full=True)

            for img_idx, img in enumerate(image_list):
                try:
                    # img là tuple: (xref, smask, width, height, bpc, colorspace, alt. colorspace, name, filter, referencer)
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]

                    images.append(
                        {
                            "data": image_bytes,
                            "ext": image_ext,
                            "xref": xref,
                            "width": img[2],
                            "height": img[3],
                        }
                    )
                except Exception as e:
                    logger.debug(
                        f"Không thể extract image {img_idx} từ trang {page_num}: {e}"
                    )
                    continue

            pages_data.append({"page_num": page_num, "text": text, "images": images})

            if not show_progress and len(pages_data) % 50 == 0:
                logger.info(
                    f"Extract text và images: {len(pages_data)}/{len(pages_to_extract)} trang"
                )

        doc.close()
        return pages_data, total

    except Exception as e:
        logger.error(f"PyMuPDF failed: {e}")
        raise


def create_docx_from_pdf(
    pdf_path: str,
    output_path: str,
    ocr_cfg: dict,
    pages: Optional[List[int]] = None,
    apply_cleanup: bool = True,
    apply_spell_check: bool = True,
) -> str:
    """
    Tạo file DOCX từ PDF có text layer, giữ lại cả text và images.
    Có thể áp dụng cleanup và spell check cho text trước khi tạo DOCX.

    Args:
        pdf_path: Đường dẫn file PDF input
        output_path: Đường dẫn file DOCX output
        ocr_cfg: Config dictionary
        pages: Danh sách số trang cần extract (1-indexed). None = tất cả trang.
        apply_cleanup: Có áp dụng AI cleanup cho text không (mặc định: True)
        apply_spell_check: Có áp dụng AI spell check cho text không (mặc định: True)

    Returns:
        str: Đường dẫn file DOCX đã tạo
    """
    if Document is None:
        raise RuntimeError(
            "python-docx chưa được cài đặt. Cài python-docx để tạo DOCX output."
        )

    logger.info(f"Tạo DOCX từ PDF: {pdf_path}")

    # Extract text và images
    pages_data, total_pages = extract_text_and_images_from_pdf(pdf_path, ocr_cfg, pages)

    if not pages_data:
        raise ValueError("Không có dữ liệu nào được extract từ PDF.")

    # Xử lý text: ghép text từ tất cả các trang và apply cleanup/spell check nếu cần
    all_text = "\n\n".join([page_info["text"] for page_info in pages_data])
    processed_text = all_text  # Mặc định là text gốc

    if apply_cleanup or apply_spell_check:
        logger.info("Đang xử lý text với AI (cleanup và spell check)...")

        # Apply cleanup nếu enabled
        cleanup_cfg = ocr_cfg.get("ai_cleanup", {})
        if apply_cleanup and cleanup_cfg.get("enabled", False):
            logger.info("🧹 Đang chạy AI Cleanup...")
            result = ai_cleanup_text(all_text, ocr_cfg)
            if isinstance(result, tuple):
                all_text, cleanup_failed_indices, cleanup_original_chunks = result
                cleanup_failed = len(cleanup_failed_indices)
                if cleanup_failed > 0:
                    logger.warning(f"AI Cleanup: {cleanup_failed} chunks failed")
            else:
                all_text = result
            logger.info("✅ Hoàn tất AI Cleanup")

        # Apply spell check nếu enabled
        spell_check_cfg = ocr_cfg.get("ai_spell_check", {})
        if apply_spell_check and spell_check_cfg.get("enabled", False):
            logger.info("✍️  Đang chạy AI Spell Check...")
            result = ai_spell_check_and_paragraph_restore(all_text, ocr_cfg)
            if isinstance(result, tuple):
                all_text, spell_check_failed_indices, spell_check_original_chunks = (
                    result
                )
                spell_check_failed = len(spell_check_failed_indices)
                if spell_check_failed > 0:
                    logger.warning(
                        f"AI Spell Check: {spell_check_failed} chunks failed"
                    )
            else:
                all_text = result
            logger.info("✅ Hoàn tất AI Spell Check")

        # Sau khi xử lý, text đã được cleanup và spell check
        # Ta sẽ lưu toàn bộ text đã xử lý và chèn vào DOCX sau tất cả images
        # Để giữ liên kết giữa images và text, ta sẽ:
        # 1. Chèn images từ tất cả các trang (giữ nguyên thứ tự)
        # 2. Chèn toàn bộ text đã xử lý sau tất cả images
        processed_text = all_text

    # Tạo DOCX document
    try:
        doc = Document()
    except Exception as e:
        raise RuntimeError(f"Không thể tạo Document object: {e}")

    # Set document properties (optional)
    try:
        doc.core_properties.title = os.path.splitext(os.path.basename(pdf_path))[0]
    except Exception:
        pass

    # python-docx tự động tạo một paragraph trống khi khởi tạo Document()
    # Ta sẽ để nó như vậy và chỉ thêm nội dung khi cần

    show_progress = bool(ocr_cfg.get("show_progress", True))
    total_items = len(pages_data)

    if show_progress and tqdm is not None:
        iterator = tqdm(pages_data, desc="Tạo DOCX", unit="trang")
    else:
        iterator = pages_data

    import io

    # Nếu có text đã được xử lý (cleanup/spell check), ta sẽ chèn text sau tất cả images
    has_processed_text = apply_cleanup or apply_spell_check

    # Bước 1: Chèn tất cả images từ tất cả các trang (giữ nguyên thứ tự)
    all_images = []
    for page_info in pages_data:
        all_images.extend(
            [(page_info["page_num"], img_info) for img_info in page_info["images"]]
        )

    # Chèn images
    images_added_count = 0
    if all_images:
        for page_num, img_info in (
            tqdm(all_images, desc="Chèn images", unit="ảnh")
            if (show_progress and tqdm and len(all_images) > 1)
            else all_images
        ):
            try:
                image_bytes = img_info.get("data")
                if not image_bytes or len(image_bytes) == 0:
                    logger.warning(f"Image data rỗng từ trang {page_num}, bỏ qua")
                    continue

                if len(image_bytes) < 10:  # Image quá nhỏ, có thể không hợp lệ
                    logger.warning(
                        f"Image từ trang {page_num} quá nhỏ ({len(image_bytes)} bytes), bỏ qua"
                    )
                    continue

                image_ext = img_info.get("ext", "png")

                # Validate image data - kiểm tra magic bytes
                is_valid = False
                if image_ext.lower() in ("jpeg", "jpg"):
                    if len(image_bytes) >= 2 and image_bytes[:2] == b"\xff\xd8":
                        is_valid = True
                elif image_ext.lower() == "png":
                    if (
                        len(image_bytes) >= 8
                        and image_bytes[:8] == b"\x89PNG\r\n\x1a\n"
                    ):
                        is_valid = True
                elif image_ext.lower() in ("gif", "bmp", "tiff", "webp"):
                    # Chấp nhận các format khác mà không validate magic bytes (để python-docx xử lý)
                    is_valid = True
                else:
                    # Format khác, thử add anyway
                    is_valid = True

                if not is_valid:
                    logger.warning(
                        f"Image từ trang {page_num} không phải {image_ext.upper()} hợp lệ, bỏ qua"
                    )
                    continue

                # Tạo image từ bytes
                img_stream = io.BytesIO(image_bytes)
                img_stream.seek(0)  # Đảm bảo stream ở đầu file

                # Thêm image vào document
                # Giới hạn kích thước để vừa với page width
                para = doc.add_paragraph()
                run = para.add_run()

                # Tính toán kích thước image (giữ tỷ lệ)
                img_width = img_info.get("width", 0)
                img_height = img_info.get("height", 0)

                try:
                    if img_width > 0 and img_height > 0:
                        # Giới hạn width tối đa là 6 inches (khoảng 15cm)
                        max_width_inches = 6.0
                        aspect_ratio = img_height / img_width

                        if img_width > 500:  # Nếu image lớn, scale xuống
                            width_inches = min(
                                max_width_inches, img_width / 96.0
                            )  # Giả định 96 DPI
                            height_inches = width_inches * aspect_ratio
                        else:
                            width_inches = img_width / 96.0
                            height_inches = img_height / 96.0

                        # Đảm bảo kích thước hợp lệ (min 0.1 inches, max 7 inches)
                        if width_inches > 0.1 and height_inches > 0.1:
                            run.add_picture(
                                img_stream,
                                width=Inches(min(width_inches, max_width_inches)),
                            )
                            images_added_count += 1
                        else:
                            run.add_picture(img_stream, width=Inches(4.0))
                            images_added_count += 1
                    else:
                        # Nếu không có thông tin size, dùng default
                        run.add_picture(img_stream, width=Inches(4.0))
                        images_added_count += 1

                    # Thêm spacing sau image
                    para_format = para.paragraph_format
                    para_format.space_after = Pt(6)
                except Exception as pic_error:
                    # Nếu add_picture thất bại, xóa paragraph trống
                    logger.warning(
                        f"Không thể add picture vào DOCX từ trang {page_num}: {pic_error}"
                    )
                    # Không cần xóa paragraph vì python-docx sẽ xử lý
                    continue

            except Exception as e:
                logger.warning(
                    f"Không thể thêm image vào DOCX từ trang {page_num}: {e}"
                )
                import traceback

                logger.debug(traceback.format_exc())
                continue

    if all_images and images_added_count == 0:
        logger.warning(
            f"Không có image nào được thêm vào DOCX (tổng {len(all_images)} images)"
        )

    # Bước 2: Chèn text đã được xử lý (nếu có) hoặc text gốc từ từng trang
    text_added = False
    if has_processed_text:
        # Chèn toàn bộ text đã được xử lý sau tất cả images
        if processed_text and processed_text.strip():
            logger.info("Đang chèn text đã được xử lý vào DOCX...")
            # Chia text thành paragraphs (dựa trên double newlines)
            paragraphs = processed_text.split("\n\n")

            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    para = doc.add_paragraph(para_text)
                    para_format = para.paragraph_format
                    para_format.space_after = Pt(6)
                    text_added = True
    else:
        # Không có xử lý → chèn text gốc từ từng trang cùng với images
        for page_info in iterator:
            page_num = page_info["page_num"]
            text = page_info["text"]

            # Thêm text của trang này
            if text and text.strip():
                # Chia text thành paragraphs (dựa trên double newlines)
                paragraphs = text.split("\n\n")

                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        para = doc.add_paragraph(para_text)
                        para_format = para.paragraph_format
                        para_format.space_after = Pt(6)
                        text_added = True

            # Thêm page break sau mỗi trang (trừ trang cuối) - chỉ khi có text hoặc images
            if page_num < pages_data[-1]["page_num"]:
                doc.add_page_break()

            if not show_progress and (page_num % 50 == 0 or page_num == total_items):
                logger.info(f"Đã xử lý {page_num}/{total_items} trang")

    # Đảm bảo document có ít nhất một paragraph hợp lệ (nếu không có gì cả)
    # Kiểm tra xem có nội dung gì không (images hoặc text)
    has_content = text_added or images_added_count > 0

    if not has_content:
        logger.warning(
            "Document không có nội dung (không có text và images). Thêm paragraph mặc định..."
        )
        try:
            # Thêm paragraph có nội dung
            # Không xóa paragraph trống vì có thể gây lỗi cấu trúc DOCX
            doc.add_paragraph("(Không có nội dung từ PDF)")
        except Exception as e:
            logger.warning(f"Không thể thêm paragraph mặc định: {e}")
            # Fallback: thử thêm vào paragraph đầu tiên nếu có
            try:
                if len(doc.paragraphs) > 0:
                    doc.paragraphs[0].text = "(Không có nội dung)"
                else:
                    doc.add_paragraph("(Không có nội dung)")
            except Exception:
                pass

    # Validate document trước khi save
    try:
        # Kiểm tra xem document có hợp lệ không
        # Document phải có ít nhất một element (paragraph với text hoặc images được embed)
        total_elements = len(doc.paragraphs)
        if total_elements == 0:
            logger.warning("Document trống, thêm paragraph mặc định")
            doc.add_paragraph("(Không có nội dung)")
        else:
            # Kiểm tra xem có paragraph nào có nội dung không (text hoặc runs - có thể chứa images)
            has_any_content = False
            for para in doc.paragraphs:
                # Kiểm tra text
                if para.text.strip():
                    has_any_content = True
                    break
                # Kiểm tra runs (có thể chứa images hoặc inline shapes)
                if len(para.runs) > 0:
                    # Nếu có runs, giả định là có nội dung (images hoặc text)
                    has_any_content = True
                    break

            if not has_any_content:
                logger.warning("Tất cả paragraphs đều trống, thêm paragraph mặc định")
                # Sử dụng paragraph đầu tiên nếu có, hoặc tạo mới
                try:
                    if len(doc.paragraphs) > 0:
                        # Thêm text vào paragraph đầu tiên
                        para = doc.paragraphs[0]
                        if not para.text.strip():
                            para.add_run("(Không có nội dung)")
                        else:
                            doc.add_paragraph("(Không có nội dung)")
                    else:
                        doc.add_paragraph("(Không có nội dung)")
                except Exception:
                    # Fallback: chỉ tạo paragraph mới
                    try:
                        doc.add_paragraph("(Không có nội dung)")
                    except Exception:
                        pass
    except Exception as e:
        logger.warning(f"Không thể validate document: {e}")

    # Lưu file
    logger.info(f"Đang lưu DOCX: {output_path}")
    try:
        doc.save(output_path)
        logger.info(f"✅ Đã tạo DOCX thành công: {output_path}")

        # Validate file sau khi save (kiểm tra file size)
        if os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            if file_size == 0:
                raise ValueError("File DOCX có kích thước 0 bytes - không hợp lệ")
            logger.info(f"📄 File size: {file_size:,} bytes")
    except Exception as e:
        logger.error(f"❌ Lỗi khi lưu DOCX: {e}")
        import traceback

        logger.error(traceback.format_exc())
        raise

    return output_path


async def _cleanup_chunk_async(
    chunk: str,
    api_key: str,
    model_name: str,
    prompt: str,
    chunk_idx: int,
    total_chunks: int,
    timeout_s: float,
    safety_settings: Optional[List[dict]] = None,
) -> str:
    """
    Cleanup một chunk text bằng AI (async).

    Args:
        safety_settings: Optional safety settings để pass vào GenerativeModel (nếu None sẽ dùng default)
    """
    # Suppress logs TRƯỚC khi import
    _suppress_google_logs()
    # Đảm bảo stderr filter đang active
    if not isinstance(sys.stderr, NoisyMessageFilter):
        original_stderr = (
            sys.stderr
            if not isinstance(sys.stderr, NoisyMessageFilter)
            else getattr(sys.stderr, "original_stream", sys.stderr)
        )
        sys.stderr = NoisyMessageFilter(original_stderr)
    import google.generativeai as genai

    genai.configure(api_key=api_key)
    # Pass safety_settings vào GenerativeModel (nếu có)
    model = (
        genai.GenerativeModel(model_name, safety_settings=safety_settings)
        if safety_settings
        else genai.GenerativeModel(model_name)
    )

    # Run trong thread pool và áp timeout để tránh treo vô hạn
    loop = asyncio.get_event_loop()
    response = await asyncio.wait_for(
        loop.run_in_executor(None, lambda: model.generate_content(prompt + chunk)),
        timeout=timeout_s,
    )

    # Kiểm tra response có hợp lệ không
    if not response or not response.candidates or len(response.candidates) == 0:
        raise ValueError(
            f"AI cleanup chunk {chunk_idx}/{total_chunks}: No candidates returned"
        )

    # Kiểm tra prompt_feedback nếu có
    if hasattr(response, "prompt_feedback") and response.prompt_feedback:
        if (
            hasattr(response.prompt_feedback, "block_reason")
            and response.prompt_feedback.block_reason
        ):
            raise ValueError(
                f"AI cleanup chunk {chunk_idx}/{total_chunks}: Blocked by safety filter: {response.prompt_feedback.block_reason}"
            )

    result = response.text.strip()
    return result


def ai_cleanup_text(text: str, ocr_cfg: dict) -> str:
    """
    Sử dụng AI để dọn rác text (header/footer, vệt đen từ scan, noise...).
    Hỗ trợ nhiều API keys để xử lý song song.
    """
    cleanup_cfg = ocr_cfg.get("ai_cleanup", {})
    cleanup_enabled = cleanup_cfg.get("enabled", False)
    if not cleanup_enabled:
        return text

    # Lấy API keys (ưu tiên từ ai_cleanup.api_keys, fallback về api_keys từ root config)
    api_keys = cleanup_cfg.get("api_keys", [])
    if not api_keys:
        # Đọc từ _root_api_keys đã lưu khi load config
        api_keys = ocr_cfg.get("_root_api_keys", [])

    if not api_keys:
        logger.warning("AI cleanup enabled nhưng không có API keys. Bỏ qua cleanup.")
        return text

    model_name = cleanup_cfg.get("model", "gemini-2.5-flash")
    max_parallel = cleanup_cfg.get("max_parallel_workers", 5)
    # Giới hạn worker theo số API keys sẵn có
    if api_keys:
        max_parallel = max(1, min(max_parallel, len(api_keys)))
    chunk_size = cleanup_cfg.get("chunk_size", 50000)
    delay = cleanup_cfg.get("delay_between_requests", 0.5)
    max_retries = cleanup_cfg.get("max_retries", 3)
    timeout_s = float(cleanup_cfg.get("ai_timeout_seconds", 120))
    show_progress = bool(ocr_cfg.get("show_progress", True))
    progress_interval = float(ocr_cfg.get("progress_log_interval_seconds", 60))

    prompt = """Bạn là một AI chuyên dọn dẹp văn bản OCR/scan. Nhiệm vụ:
1. Loại bỏ header/footer lặp lại ở mỗi trang
2. Loại bỏ các ký tự rác, vệt đen vô nghĩa từ quá trình scan
3. Loại bỏ số trang, watermark
4. Giữ nguyên nội dung chính của văn bản
5. Chuẩn hóa khoảng trắng thừa
6. Giữ nguyên định dạng đoạn văn

Trả về chỉ văn bản đã được dọn dẹp, không giải thích thêm.

Văn bản cần dọn dẹp:
"""

    try:
        # Chia nhỏ text nếu quá dài
        if len(text) <= chunk_size:
            # Text ngắn, xử lý trực tiếp
            logger.info("AI Cleanup: Text ngắn, xử lý trực tiếp (1 chunk)")
            # Suppress logs TRƯỚC khi import
            _suppress_google_logs()
            # Đảm bảo stderr filter đang active
            if not isinstance(sys.stderr, NoisyMessageFilter):
                original_stderr = (
                    sys.stderr
                    if not isinstance(sys.stderr, NoisyMessageFilter)
                    else getattr(sys.stderr, "original_stream", sys.stderr)
                )
                sys.stderr = NoisyMessageFilter(original_stderr)
            # Build safety settings từ config
            safety_level = ocr_cfg.get("safety_level", "BLOCK_ONLY_HIGH")
            safety_settings = _build_safety_settings(safety_level)

            import google.generativeai as genai

            genai.configure(api_key=api_keys[0])
            model = genai.GenerativeModel(model_name, safety_settings=safety_settings)
            response = model.generate_content(prompt + text)

            # Kiểm tra nếu response bị block (mặc dù đã set BLOCK_NONE, nhưng vẫn check để an toàn)
            if hasattr(response, "prompt_feedback") and response.prompt_feedback:
                block_reason = getattr(response.prompt_feedback, "block_reason", None)
                if block_reason:
                    logger.warning(
                        f"AI Cleanup bị block: {block_reason}. Sử dụng text gốc."
                    )
                    return (text, [0], [text])  # Return text gốc với failed index

            if not hasattr(response, "candidates") or not response.candidates:
                logger.warning("AI Cleanup không có candidates. Sử dụng text gốc.")
                return (text, [0], [text])  # Return text gốc với failed index

            logger.info(
                "AI Cleanup: Hoàn tất. Thành công: 1/1 chunk, Thất bại: 0/1 chunk."
            )
            cleaned_text = response.text.strip()
            return (
                cleaned_text,
                [],
                [text],
            )  # (result_text, failed_indices, original_chunks)

        # Build safety settings từ config
        safety_level = ocr_cfg.get("safety_level", "BLOCK_ONLY_HIGH")
        safety_settings = _build_safety_settings(safety_level)

        # Text dài, chia nhỏ ở ranh giới câu và xử lý song song
        text_chunks = _split_text_at_sentence_boundaries(text, chunk_size)
        total_chunks = len(text_chunks)
        logger.info(
            f"AI Cleanup: Chia thành {total_chunks} chunks (ở ranh giới câu), xử lý song song với {len(api_keys)} API keys"
        )
        logger.info(f"AI Cleanup: Safety level: {safety_level}")
        logger.info("AI Cleanup: Bắt đầu xử lý...")

        # Chạy async cleanup với safety settings
        result_text, success_count, failure_count, failed_indices = asyncio.run(
            _ai_cleanup_parallel(
                text_chunks,
                api_keys,
                model_name,
                prompt,
                max_parallel,
                delay,
                show_progress,
                timeout_s,
                max_retries,
                progress_interval,
                safety_settings,
            )
        )
        logger.info(
            f"AI Cleanup: Hoàn tất. Thành công: {success_count}/{total_chunks} chunks, Thất bại: {failure_count}/{total_chunks} chunks (đã lưu nội dung gốc)."
        )

        # Tự động retry các chunks failed sau khi hoàn tất tất cả chunks khác
        if failure_count > 0:
            auto_retry = cleanup_cfg.get("auto_retry_failed", True)  # Mặc định: true
            if auto_retry:
                logger.info(
                    f"AI Cleanup: Tự động retry {failure_count} chunks failed..."
                )
                retry_results, still_failed = _retry_failed_chunks_cleanup(
                    failed_indices, text_chunks, api_keys, model_name, prompt, ocr_cfg
                )

                # Merge lại text từ retry results
                if retry_results:
                    cleanup_chunks_list = list(text_chunks)
                    for idx, retry_text in retry_results.items():
                        if idx < len(cleanup_chunks_list):
                            cleanup_chunks_list[idx] = retry_text
                    result_text = "\n\n".join(cleanup_chunks_list)

                    retry_success = len(retry_results) - len(still_failed)
                    logger.info(
                        f"AI Cleanup Auto Retry: {retry_success}/{failure_count} chunks retry thành công."
                    )
                    if still_failed:
                        logger.warning(
                            f"AI Cleanup Auto Retry: {len(still_failed)} chunks vẫn failed sau retry."
                        )
                        # Cập nhật failed_indices với still_failed
                        failed_indices = still_failed
                    else:
                        logger.info(
                            "AI Cleanup Auto Retry: Tất cả chunks failed đã được retry thành công!"
                        )
                        failed_indices = []  # Tất cả đã thành công
                else:
                    logger.warning("AI Cleanup Auto Retry: Không có kết quả retry.")

        # Trả về text đã merge, failed_indices, và toàn bộ chunks (để có thể rebuild sau retry)
        return (result_text, failed_indices, text_chunks)

    except Exception as e:
        logger.error(f"AI cleanup failed: {e}. Trả về text gốc.")
        return (text, [], [])  # Trả về tuple nhất quán


async def _ai_cleanup_parallel(
    text_chunks: List[str],
    api_keys: List[str],
    model_name: str,
    prompt: str,
    max_parallel: int,
    delay: float,
    show_progress: bool,
    timeout_s: float,
    max_retries: int,
    progress_interval: float,
    safety_settings: Optional[List[dict]] = None,
) -> tuple[str, int, int, List[int]]:
    """
    Xử lý song song nhiều chunks với nhiều API keys.

    Args:
        safety_settings: Optional safety settings để pass vào GenerativeModel
    """
    # Tạo queue cho API keys
    key_queue = asyncio.Queue()
    for key in api_keys:
        await key_queue.put(key)

    cleaned_chunks: List[tuple[int, str]] = []  # (index, cleaned_text)
    semaphore = asyncio.Semaphore(max_parallel)
    total = len(text_chunks)
    failures = 0
    failed_indices: List[int] = []

    async def process_chunk(chunk: str, chunk_idx: int) -> tuple[int, str]:
        nonlocal failures, failed_indices  # Khai báo nonlocal ở đầu function
        async with semaphore:
            retries = 0
            api_key = None
            last_error = None
            while retries < max_retries:
                try:
                    api_key = await key_queue.get()
                    cleaned = await _cleanup_chunk_async(
                        chunk,
                        api_key,
                        model_name,
                        prompt,
                        chunk_idx,
                        len(text_chunks),
                        timeout_s,
                        safety_settings,
                    )
                    # Thành công - return ngay
                    return (chunk_idx, cleaned)
                except Exception as e:
                    last_error = e
                    retries += 1
                    if retries < max_retries:
                        logger.debug(
                            f"AI cleanup chunk {chunk_idx} failed (attempt {retries}/{max_retries}): {type(e).__name__}: {e}. Retrying..."
                        )
                        await asyncio.sleep(delay * retries)
                    else:
                        # Đã retry hết
                        failures += 1
                        failed_indices.append(chunk_idx)
                        logger.warning(
                            f"AI cleanup chunk {chunk_idx} failed after {max_retries} retries with {type(e).__name__}: {e}"
                        )
                        return (chunk_idx, chunk)  # Trả về chunk gốc
                finally:
                    if api_key:
                        try:
                            await key_queue.put(
                                api_key
                            )  # Trả key về queue dù thành công hay lỗi
                        except Exception:
                            pass
                    await asyncio.sleep(delay)

            # Nếu đến đây (không nên xảy ra)
            failures += 1
            failed_indices.append(chunk_idx)
            logger.warning(
                f"AI cleanup chunk {chunk_idx} failed after all retries. Last error: {last_error}"
            )
            return (chunk_idx, chunk)

    # Tạo tasks cho tất cả chunks
    tasks = [process_chunk(chunk, idx) for idx, chunk in enumerate(text_chunks)]

    # Xử lý và log tiến độ định kỳ
    results = []
    if show_progress:
        start_ts = time.time()
        last_log = start_ts
        completed = 0
        async for result in _as_completed_iter(tasks):
            results.append(result)
            completed += 1
            now = time.time()
            if (now - last_log) >= max(5.0, progress_interval):
                elapsed = now - start_ts
                avg = elapsed / completed if completed > 0 else 0.0
                remaining = max(len(tasks) - completed, 0) * avg
                logger.info(
                    f"AI Cleanup: {completed}/{len(tasks)} chunks • TB {avg:.2f}s/chunk • ETA ~{remaining:.0f}s"
                )
                last_log = now
    else:
        results = await asyncio.gather(*tasks)
    cleaned_chunks = sorted(results, key=lambda x: x[0])
    success_count = total - failures
    if failures > 0:
        logger.warning(
            f"AI Cleanup: {failures}/{total} chunks failed. Tiếp tục với nội dung gốc cho các chunk lỗi."
        )

    # Ghép các chunks theo thứ tự
    result_text = "\n\n".join([text for _, text in cleaned_chunks])
    return (result_text, success_count, failures, failed_indices)


async def _as_completed_iter(coros):
    for fut in asyncio.as_completed(coros):
        yield await fut


def _split_text_at_sentence_boundaries(text: str, max_chunk_size: int) -> List[str]:
    """
    Chia text thành chunks ở ranh giới câu (kết thúc bằng dấu chấm câu).
    Tham khảo thuật toán từ SmartChunker._split_long_paragraph để đảm bảo không cắt giữa câu.

    Args:
        text: Văn bản cần chia
        max_chunk_size: Kích thước tối đa của mỗi chunk (tính theo ký tự)

    Returns:
        List[str]: Danh sách các chunks đã được chia ở ranh giới câu
    """
    import re

    if not text or len(text) <= max_chunk_size:
        return [text] if text else []

    # Pattern để tìm ranh giới câu: . ! ? (cả tiếng Anh) và 。！？ (tiếng Trung)
    # Hỗ trợ các dấu ngoặc kép có thể đi kèm: ["']? (cho tiếng Anh) và » (cho một số ngôn ngữ)
    sentence_pattern = re.compile(r'([.!?。！？]["\'»]?\s*)')

    # Tìm tất cả các vị trí kết thúc câu
    parts = sentence_pattern.split(text)

    # Ghép lại các phần để tạo sentences (mỗi sentence bao gồm nội dung + dấu câu)
    sentences = []
    for i in range(0, len(parts) - 1, 2):
        if i + 1 < len(parts):
            sentence = (parts[i] + parts[i + 1]).strip()
            if sentence:
                sentences.append(sentence)

    # Xử lý phần cuối cùng nếu không kết thúc bằng dấu câu
    if len(parts) % 2 == 1 and parts[-1].strip():
        sentences.append(parts[-1].strip())

    # Lọc bỏ các câu rỗng
    sentences = [sent for sent in sentences if sent.strip()]

    if not sentences:
        return [text]

    # Gom các sentences thành chunks, đảm bảo không vượt quá max_chunk_size
    chunks = []
    current_chunk = []
    current_size = 0

    for sentence in sentences:
        sent_size = len(sentence)

        # Nếu sentence đơn lẻ quá dài, phải cắt (trường hợp hiếm)
        if sent_size > max_chunk_size:
            # Nếu đang có chunk tích lũy, lưu nó trước
            if current_chunk:
                chunks.append(" ".join(current_chunk))
                current_chunk = []
                current_size = 0

            # Chia sentence dài thành nhiều phần nhỏ hơn
            # Ưu tiên cắt ở khoảng trắng nếu có thể
            words = sentence.split()
            temp_chunk = []
            temp_size = 0

            for word in words:
                word_size = len(word) + 1  # +1 cho space
                if temp_size + word_size > max_chunk_size and temp_chunk:
                    # Lưu chunk hiện tại
                    chunks.append(" ".join(temp_chunk))
                    temp_chunk = [word]
                    temp_size = len(word)
                else:
                    temp_chunk.append(word)
                    temp_size += word_size

            if temp_chunk:
                chunks.append(" ".join(temp_chunk))
        else:
            # Kiểm tra nếu thêm sentence này vào chunk hiện tại có vượt quá max_chunk_size không
            # Nếu đã có sentences trong chunk, cần thêm 1 ký tự cho space khi join
            space_needed = 1 if current_chunk else 0
            if (
                current_size + sent_size + space_needed > max_chunk_size
                and current_chunk
            ):
                # Lưu chunk hiện tại và bắt đầu chunk mới
                chunks.append(" ".join(current_chunk))
                current_chunk = [sentence]
                current_size = sent_size
            else:
                # Thêm sentence vào chunk hiện tại
                current_chunk.append(sentence)
                current_size += sent_size + space_needed

    # Lưu chunk cuối cùng nếu có
    if current_chunk:
        chunks.append(" ".join(current_chunk))

    # Nếu không chia được gì (trường hợp hiếm), trả về toàn bộ text
    if not chunks:
        return [text]

    return chunks


def _preprocess_line_breaks(text: str) -> str:
    """
    Preprocessing: Nối lại các câu bị ngắt do line breaks khi convert PDF → TXT.
    Chỉ xử lý các trường hợp rõ ràng, các trường hợp phức tạp sẽ để AI xử lý.
    """
    import re

    lines = text.split("\n")
    if not lines:
        return text

    result_lines = []
    i = 0

    while i < len(lines):
        current_line = lines[i].strip()

        # Nếu dòng rỗng → giữ nguyên (đây là paragraph break)
        if not current_line:
            result_lines.append("")
            i += 1
            continue

        # Bắt đầu từ dòng hiện tại, cố gắng nối các dòng tiếp theo nếu thỏa điều kiện
        merged_line = current_line

        # Kiểm tra và nối các dòng tiếp theo liên tục
        while i + 1 < len(lines):
            next_line = lines[i + 1].strip()

            # Nếu dòng tiếp theo rỗng
            if not next_line:
                # Kiểm tra xem có phải paragraph break thực sự không
                # Nếu dòng hiện tại kết thúc bằng dấu câu → paragraph break thực sự
                if re.search(r"[.!?]$", merged_line):
                    break

                # Nếu không, có thể là line break do PDF format, kiểm tra dòng sau nữa
                if i + 2 < len(lines):
                    next_next_line = lines[i + 2].strip()
                    if next_next_line:
                        # Kiểm tra dòng sau dòng trống
                        first_char = next_next_line[0]
                        next_starts_with_upper = (
                            first_char.isupper() and first_char.isalpha()
                        )
                        next_starts_with_number = bool(
                            re.match(r"^\d+", next_next_line)
                        )
                        next_starts_with_bullet = bool(
                            re.match(r"^[•·\-*]\s", next_next_line)
                        )

                        # Nếu dòng sau dòng trống bắt đầu bằng chữ hoa/số/bullet → paragraph break thực sự
                        if (
                            next_starts_with_upper
                            or next_starts_with_number
                            or next_starts_with_bullet
                        ):
                            break

                        # Nếu dòng sau dòng trống bắt đầu bằng chữ thường → có thể là câu bị ngắt
                        # Bỏ qua dòng trống và tiếp tục với dòng sau
                        next_line = next_next_line
                        i += 1  # Skip dòng trống
                    else:
                        # Không còn dòng nào → dừng
                        break
                else:
                    # Không còn dòng nào → dừng
                    break

            # Dòng hiện tại (đã merged) KHÔNG kết thúc bằng dấu kết thúc câu
            ends_with_punctuation = bool(re.search(r"[.!?]$", merged_line))

            if not ends_with_punctuation:
                # Kiểm tra nếu dòng tiếp theo bắt đầu bằng chữ hoa
                # Dùng phương pháp đơn giản: kiểm tra ký tự đầu tiên có phải chữ hoa không
                if next_line:
                    first_char = next_line[0]
                    # Kiểm tra nếu là chữ cái và viết hoa (hỗ trợ Unicode)
                    next_starts_with_upper = (
                        first_char.isupper() and first_char.isalpha()
                    )
                else:
                    next_starts_with_upper = False

                next_starts_with_number = bool(re.match(r"^\d+", next_line))
                next_starts_with_bullet = bool(re.match(r"^[•·\-*]\s", next_line))

                # Nếu dòng tiếp theo KHÔNG bắt đầu bằng chữ hoa VÀ không phải số/bullet
                # → Có thể là câu bị ngắt, nối lại
                if (
                    not next_starts_with_upper
                    and not next_starts_with_number
                    and not next_starts_with_bullet
                ):
                    # Nối với dòng tiếp theo
                    merged_line = merged_line.rstrip() + " " + next_line.lstrip()
                    i += 1
                    # Tiếp tục kiểm tra dòng tiếp theo
                    continue

            # Không thỏa điều kiện nối → dừng
            break

        # Lưu dòng đã merged (hoặc dòng gốc nếu không merge)
        result_lines.append(merged_line)
        i += 1

    return "\n".join(result_lines)


async def _spell_check_chunk_async(
    chunk: str,
    api_key: str,
    model_name: str,
    prompt: str,
    chunk_idx: int,
    total_chunks: int,
    timeout_s: float,
    safety_settings: Optional[List[dict]] = None,
) -> str:
    """
    Soát lỗi chính tả và phục hồi paragraph cho một chunk text bằng AI (async).

    Args:
        safety_settings: Optional safety settings để pass vào GenerativeModel (nếu None sẽ dùng default)
    """
    # Suppress logs TRƯỚC khi import
    _suppress_google_logs()
    # Đảm bảo stderr filter đang active
    if not isinstance(sys.stderr, NoisyMessageFilter):
        original_stderr = (
            sys.stderr
            if not isinstance(sys.stderr, NoisyMessageFilter)
            else getattr(sys.stderr, "original_stream", sys.stderr)
        )
        sys.stderr = NoisyMessageFilter(original_stderr)
    import google.generativeai as genai

    genai.configure(api_key=api_key)
    # Pass safety_settings vào GenerativeModel (nếu có)
    model = (
        genai.GenerativeModel(model_name, safety_settings=safety_settings)
        if safety_settings
        else genai.GenerativeModel(model_name)
    )

    # Run trong thread pool và áp timeout để tránh treo vô hạn
    loop = asyncio.get_event_loop()
    response = await asyncio.wait_for(
        loop.run_in_executor(None, lambda: model.generate_content(prompt + chunk)),
        timeout=timeout_s,
    )
    # Kiểm tra response có hợp lệ không
    if not response or not response.candidates or len(response.candidates) == 0:
        raise ValueError(
            f"AI spell check chunk {chunk_idx}/{total_chunks}: No candidates returned"
        )

    # Kiểm tra prompt_feedback nếu có
    if hasattr(response, "prompt_feedback") and response.prompt_feedback:
        if (
            hasattr(response.prompt_feedback, "block_reason")
            and response.prompt_feedback.block_reason
        ):
            raise ValueError(
                f"AI spell check chunk {chunk_idx}/{total_chunks}: Blocked by safety filter: {response.prompt_feedback.block_reason}"
            )

    result = response.text.strip()
    return result


async def _ai_spell_check_parallel(
    text_chunks: List[str],
    api_keys: List[str],
    model_name: str,
    prompt: str,
    max_parallel: int,
    delay: float,
    show_progress: bool,
    timeout_s: float,
    max_retries: int,
    progress_interval: float,
    safety_settings: Optional[List[dict]] = None,
) -> tuple[str, int, int, List[int]]:
    """
    Xử lý song song nhiều chunks với nhiều API keys cho spell check.

    Args:
        safety_settings: Optional safety settings để pass vào GenerativeModel
    """
    # Tạo queue cho API keys
    key_queue = asyncio.Queue()
    for key in api_keys:
        await key_queue.put(key)

    processed_chunks: List[tuple[int, str]] = []  # (index, processed_text)
    semaphore = asyncio.Semaphore(max_parallel)
    total = len(text_chunks)
    failures = 0
    failed_indices: List[int] = []

    async def process_chunk(chunk: str, chunk_idx: int) -> tuple[int, str]:
        nonlocal failures, failed_indices  # Khai báo nonlocal ở đầu function
        async with semaphore:
            retries = 0
            api_key = None
            last_error = None
            while retries < max_retries:
                try:
                    api_key = await key_queue.get()
                    processed = await _spell_check_chunk_async(
                        chunk,
                        api_key,
                        model_name,
                        prompt,
                        chunk_idx,
                        len(text_chunks),
                        timeout_s,
                        safety_settings,
                    )
                    # Thành công - return ngay
                    return (chunk_idx, processed)
                except Exception as e:
                    last_error = e
                    retries += 1
                    if retries < max_retries:
                        logger.debug(
                            f"AI spell check chunk {chunk_idx} failed (attempt {retries}/{max_retries}): {type(e).__name__}: {e}. Retrying..."
                        )
                        await asyncio.sleep(delay * retries)
                    else:
                        # Đã retry hết
                        failures += 1
                        failed_indices.append(chunk_idx)
                        logger.warning(
                            f"AI spell check chunk {chunk_idx} failed after {max_retries} retries with {type(e).__name__}: {e}"
                        )
                        return (chunk_idx, chunk)  # Trả về chunk gốc
                finally:
                    if api_key:
                        try:
                            await key_queue.put(
                                api_key
                            )  # Trả key về queue dù thành công hay lỗi
                        except Exception:
                            pass
                    await asyncio.sleep(delay)

            # Nếu đến đây (không nên xảy ra)
            failures += 1
            failed_indices.append(chunk_idx)
            logger.warning(
                f"AI spell check chunk {chunk_idx} failed after all retries. Last error: {last_error}"
            )
            return (chunk_idx, chunk)

    # Tạo tasks cho tất cả chunks
    tasks = [process_chunk(chunk, idx) for idx, chunk in enumerate(text_chunks)]

    # Xử lý và log tiến độ định kỳ
    results = []
    if show_progress:
        start_ts = time.time()
        last_log = start_ts
        completed = 0
        async for result in _as_completed_iter(tasks):
            results.append(result)
            completed += 1
            now = time.time()
            if (now - last_log) >= max(5.0, progress_interval):
                elapsed = now - start_ts
                avg = elapsed / completed if completed > 0 else 0.0
                remaining = max(len(tasks) - completed, 0) * avg
                logger.info(
                    f"AI Spell Check: {completed}/{len(tasks)} chunks • TB {avg:.2f}s/chunk • ETA ~{remaining:.0f}s"
                )
                last_log = now
    else:
        results = await asyncio.gather(*tasks)
    processed_chunks = sorted(results, key=lambda x: x[0])
    success_count = total - failures

    if failures > 0:
        logger.warning(
            f"AI Spell Check: {failures}/{total} chunks failed. Tiếp tục với nội dung gốc cho các chunk lỗi."
        )
    else:
        logger.info(f"AI Spell Check: Tất cả {total} chunks đã được xử lý thành công.")

    # Ghép các chunks theo thứ tự
    result_text = "\n\n".join([text for _, text in processed_chunks])
    return (result_text, success_count, failures, failed_indices)


def ai_spell_check_and_paragraph_restore(text: str, ocr_cfg: dict) -> str:
    """
    Sử dụng AI để soát lỗi chính tả và phục hồi cấu trúc paragraph.
    Đặc biệt chú ý bảo vệ toàn vẹn nội dung (không thay đổi ý nghĩa).
    Hỗ trợ nhiều API keys để xử lý song song.
    """
    spell_check_cfg = ocr_cfg.get("ai_spell_check", {})
    spell_check_enabled = spell_check_cfg.get("enabled", False)
    if not spell_check_enabled:
        return text

    # Ghi chú: Không dùng preprocessing rule-based vì AI sẽ phân tích ngữ cảnh tốt hơn
    # Hàm _preprocess_line_breaks vẫn được giữ lại nếu cần dùng trong tương lai
    # text = _preprocess_line_breaks(text)  # Tạm tắt để AI làm toàn bộ

    # Lấy API keys (ưu tiên từ ai_spell_check.api_keys, fallback về api_keys từ root config)
    api_keys = spell_check_cfg.get("api_keys", [])
    if not api_keys:
        # Đọc từ _root_api_keys đã lưu khi load config
        api_keys = ocr_cfg.get("_root_api_keys", [])

    if not api_keys:
        logger.warning(
            "AI spell check enabled nhưng không có API keys. Bỏ qua spell check."
        )
        return text

    model_name = spell_check_cfg.get("model", "gemini-2.5-flash")
    max_parallel = spell_check_cfg.get("max_parallel_workers", 5)
    # Giới hạn worker theo số API keys sẵn có
    if api_keys:
        max_parallel = max(1, min(max_parallel, len(api_keys)))
    chunk_size = spell_check_cfg.get("chunk_size", 50000)
    delay = spell_check_cfg.get("delay_between_requests", 0.5)
    max_retries = spell_check_cfg.get("max_retries", 3)
    timeout_s = float(spell_check_cfg.get("ai_timeout_seconds", 120))
    show_progress = bool(ocr_cfg.get("show_progress", True))
    progress_interval = float(ocr_cfg.get("progress_log_interval_seconds", 60))

    prompt = """Bạn là một AI chuyên soát lỗi chính tả và phục hồi cấu trúc văn bản OCR. Nhiệm vụ chính của bạn là PHÂN TÍCH NGỮ CẢNH và QUYẾT ĐỊNH THÔNG MINH.

=== NHIỆM VỤ CHÍNH: PHÂN TÍCH VÀ PHỤC HỒI CÂU BỊ NGẮT (Ưu tiên cao nhất) ===

Bạn cần ĐỌC KỸ NỘI DUNG và PHÂN TÍCH để phân biệt:

A. CÂU BỊ NGẮT DO CONVERT PDF → TXT (CẦN NỐI LẠI):
   - Đọc ngữ cảnh: Nếu dòng trước chưa hoàn thành ý và dòng sau tiếp nối ý đó → nối lại
   - Ví dụ: 
     * "Our client is also the owner of Vietnam Trade Mark Registration No. 315843 for "MICROBAN"
       in Class 5 covering..." 
     → Phân tích: "in Class 5" tiếp nối câu trước → NỐI LẠI thành một câu
   
   - Dấu hiệu cần nối:
     * Dòng trước không kết thúc bằng dấu câu (. ! ?) HOẶC kết thúc bằng dấu phẩy, hai chấm
     * Dòng sau bắt đầu bằng chữ thường (tiếp nối câu trước)
     * Nội dung dòng sau về mặt ngữ pháp và ngữ nghĩa là phần tiếp theo của câu trước
     * Đọc toàn bộ ngữ cảnh để hiểu rõ mối quan hệ

B. NGẮT PARAGRAPH CÓ CHỦ ĐÍCH (KHÔNG NỐI):
   - Đọc ngữ cảnh: Nếu dòng sau là ý mới, chủ đề mới, hoặc đoạn văn mới → KHÔNG nối
   - Ví dụ:
     * "...attached as Exhibit 1.
       
       Khách hàng của chúng tôi là chủ sở hữu..."
     → Phân tích: Đây là đoạn mới (chuyển từ tiếng Anh sang tiếng Việt) → KHÔNG NỐI
   
   - Dấu hiệu KHÔNG nối:
     * Dòng trước kết thúc bằng dấu chấm (. ! ?) và dòng sau bắt đầu bằng chữ hoa
     * Dòng sau là câu đầu tiên của một đoạn mới (ý tưởng mới, chủ đề mới)
     * Có sự thay đổi rõ ràng về ngữ cảnh (ví dụ: chuyển từ phần này sang phần khác)
     * Đọc toàn bộ ngữ cảnh để xác định đây là ngắt đoạn có chủ đích

QUY TRÌNH PHÂN TÍCH:
1. ĐỌC toàn bộ văn bản để hiểu cấu trúc và ngữ cảnh
2. PHÂN TÍCH từng vị trí ngắt dòng:
   - Xem xét nội dung trước và sau dòng ngắt
   - Đánh giá mối quan hệ ngữ pháp và ngữ nghĩa
   - Xác định đây là câu bị ngắt hay ngắt đoạn có chủ đích
3. QUYẾT ĐỊNH:
   - Nếu là câu bị ngắt → NỐI lại (thay line break bằng space)
   - Nếu là ngắt đoạn có chủ đích → GIỮ NGUYÊN (có thể thêm dòng trống nếu cần)
4. ÁP DỤNG nhất quán cho toàn bộ văn bản

=== CÁC NHIỆM VỤ KHÁC ===

1. SOÁT LỖI CHÍNH TẢ:
   - Sửa các lỗi chính tả do OCR (ví dụ: "Kíng" → "Kính", "hang" → "hàng")
   - Sửa các lỗi chính tả thông thường
   - KHÔNG thay đổi từ ngữ chuyên ngành, tên riêng, địa danh
   - KHÔNG thay đổi số liệu, ngày tháng, địa chỉ

2. PHỤC HỒI CẤU TRÚC PARAGRAPH:
   - Sau khi đã nối các câu bị ngắt, xác định các ngắt đoạn hợp lý
   - Mỗi đoạn văn nên có một ý chính hoàn chỉnh
   - Giữ nguyên các dòng trống giữa các đoạn đã được xác định là có chủ đích
   - Đảm bảo các câu trong một đoạn có liên quan với nhau

3. BẢO VỆ TOÀN VẸN NỘI DUNG:
   - TUYỆT ĐỐI KHÔNG thay đổi ý nghĩa của văn bản
   - KHÔNG thêm, bớt, hoặc diễn giải lại nội dung
   - KHÔNG thay đổi thứ tự từ trong câu (chỉ nối lại khi cần)
   - GIỮ NGUYÊN định dạng đặc biệt (bullet points, numbered lists, bảng)
   - GIỮ NGUYÊN các từ viết hoa nếu chúng là tên riêng, thuật ngữ

4. ĐỊNH DẠNG:
   - Giữ nguyên định dạng văn bản song ngữ (nếu có)
   - Giữ nguyên các dấu câu quan trọng
   - Chuẩn hóa khoảng trắng thừa giữa các từ (nhưng không thay đổi paragraph breaks hợp lý)
   - Đảm bảo mỗi câu kết thúc bằng dấu câu thích hợp

=== NGUYÊN TẮC QUAN TRỌNG ===

- SỬ DỤNG SỨC MẠNH PHÂN TÍCH NGỮ CẢNH: Đọc và hiểu nội dung, không chỉ dựa vào quy tắc cú pháp
- QUYẾT ĐỊNH THÔNG MINH: Mỗi quyết định nối hay không nối phải dựa trên phân tích ngữ cảnh cụ thể
- NHẤT QUÁN: Áp dụng cùng một tiêu chuẩn phân tích cho toàn bộ văn bản
- BẢO TOÀN Ý NGHĨA: Chỉ điều chỉnh cấu trúc, KHÔNG thay đổi nội dung hoặc ý nghĩa

Trả về chỉ văn bản đã được soát và phục hồi, không giải thích thêm.

Văn bản cần phân tích và xử lý:
"""

    try:
        # Chia nhỏ text nếu quá dài
        if len(text) <= chunk_size:
            # Text ngắn, xử lý trực tiếp
            logger.info("AI Spell Check: Text ngắn, xử lý trực tiếp (1 chunk)")
            # Suppress logs TRƯỚC khi import
            _suppress_google_logs()
            # Đảm bảo stderr filter đang active
            if not isinstance(sys.stderr, NoisyMessageFilter):
                original_stderr = (
                    sys.stderr
                    if not isinstance(sys.stderr, NoisyMessageFilter)
                    else getattr(sys.stderr, "original_stream", sys.stderr)
                )
                sys.stderr = NoisyMessageFilter(original_stderr)
            # Build safety settings từ config
            safety_level = ocr_cfg.get("safety_level", "BLOCK_ONLY_HIGH")
            safety_settings = _build_safety_settings(safety_level)

            import google.generativeai as genai

            genai.configure(api_key=api_keys[0])
            model = genai.GenerativeModel(model_name, safety_settings=safety_settings)
            response = model.generate_content(prompt + text)

            # Kiểm tra nếu response bị block (mặc dù đã set BLOCK_NONE, nhưng vẫn check để an toàn)
            if hasattr(response, "prompt_feedback") and response.prompt_feedback:
                block_reason = getattr(response.prompt_feedback, "block_reason", None)
                if block_reason:
                    logger.warning(
                        f"AI Spell Check bị block: {block_reason}. Sử dụng text gốc."
                    )
                    return (text, [0], [text])  # Return text gốc với failed index

            if not hasattr(response, "candidates") or not response.candidates:
                logger.warning("AI Spell Check không có candidates. Sử dụng text gốc.")
                return (text, [0], [text])  # Return text gốc với failed index

            logger.info(
                "AI Spell Check: Hoàn tất. Thành công: 1/1 chunk, Thất bại: 0/1 chunk."
            )
            checked_text = response.text.strip()
            return (
                checked_text,
                [],
                [text],
            )  # (result_text, failed_indices, original_chunks)

        # Build safety settings từ config
        safety_level = ocr_cfg.get("safety_level", "BLOCK_ONLY_HIGH")
        safety_settings = _build_safety_settings(safety_level)

        # Text dài, chia nhỏ ở ranh giới câu và xử lý song song
        text_chunks = _split_text_at_sentence_boundaries(text, chunk_size)
        total_chunks = len(text_chunks)
        logger.info(
            f"AI Spell Check: Chia thành {total_chunks} chunks (ở ranh giới câu), xử lý song song với {len(api_keys)} API keys"
        )
        logger.info(f"AI Spell Check: Safety level: {safety_level}")
        logger.info("AI Spell Check: Bắt đầu xử lý...")

        # Chạy async spell check với safety settings
        result_text, success_count, failure_count, failed_indices = asyncio.run(
            _ai_spell_check_parallel(
                text_chunks,
                api_keys,
                model_name,
                prompt,
                max_parallel,
                delay,
                show_progress,
                timeout_s,
                max_retries,
                progress_interval,
                safety_settings,
            )
        )
        logger.info(
            f"AI Spell Check: Hoàn tất. Thành công: {success_count}/{total_chunks} chunks, Thất bại: {failure_count}/{total_chunks} chunks (đã lưu nội dung gốc)."
        )

        # Tự động retry các chunks failed sau khi hoàn tất tất cả chunks khác
        if failure_count > 0:
            auto_retry = spell_check_cfg.get(
                "auto_retry_failed", True
            )  # Mặc định: true
            if auto_retry:
                logger.info(
                    f"AI Spell Check: Tự động retry {failure_count} chunks failed..."
                )
                retry_results, still_failed = _retry_failed_chunks_spell_check(
                    failed_indices, text_chunks, api_keys, model_name, prompt, ocr_cfg
                )

                # Merge lại text từ retry results
                if retry_results:
                    spell_check_chunks_list = list(text_chunks)
                    for idx, retry_text in retry_results.items():
                        if idx < len(spell_check_chunks_list):
                            spell_check_chunks_list[idx] = retry_text
                    result_text = "\n\n".join(spell_check_chunks_list)

                    retry_success = len(retry_results) - len(still_failed)
                    logger.info(
                        f"AI Spell Check Auto Retry: {retry_success}/{failure_count} chunks retry thành công."
                    )
                    if still_failed:
                        logger.warning(
                            f"AI Spell Check Auto Retry: {len(still_failed)} chunks vẫn failed sau retry."
                        )
                        # Cập nhật failed_indices với still_failed
                        failed_indices = still_failed
                    else:
                        logger.info(
                            "AI Spell Check Auto Retry: Tất cả chunks failed đã được retry thành công!"
                        )
                        failed_indices = []  # Tất cả đã thành công
                else:
                    logger.warning("AI Spell Check Auto Retry: Không có kết quả retry.")

        # Trả về text đã merge, failed_indices, và toàn bộ chunks (để có thể rebuild sau retry)
        return (result_text, failed_indices, text_chunks)

    except Exception as e:
        logger.error(f"AI spell check failed: {e}. Trả về text gốc.")
        return (text, [], [])  # Trả về tuple nhất quán


def _retry_failed_chunks_cleanup(
    failed_indices: List[int],
    all_chunks: List[str],
    api_keys: List[str],
    model_name: str,
    prompt: str,
    ocr_cfg: dict,
) -> tuple[dict[int, str], List[int]]:
    """Retry các chunk failed cho AI Cleanup. Trả về dict {idx: processed_text} và danh sách still_failed."""
    if not failed_indices or not all_chunks:
        return ({}, [])

    cleanup_cfg = ocr_cfg.get("ai_cleanup", {})
    timeout_s = float(cleanup_cfg.get("ai_timeout_seconds", 240))

    failed_chunks = [
        (idx, all_chunks[idx]) for idx in failed_indices if idx < len(all_chunks)
    ]
    logger.info(f"AI Cleanup Retry: Đang retry {len(failed_chunks)} chunks failed...")

    # Build safety settings từ config
    safety_level = ocr_cfg.get("safety_level", "BLOCK_ONLY_HIGH")
    safety_settings = _build_safety_settings(safety_level)

    async def _retry_chunk(idx: int, chunk: str) -> tuple[int, str]:
        for key in api_keys:
            try:
                result = await _cleanup_chunk_async(
                    chunk,
                    key,
                    model_name,
                    prompt,
                    idx,
                    len(failed_chunks),
                    timeout_s,
                    safety_settings,
                )
                return (idx, result)
            except Exception:
                continue
        return (idx, chunk)  # Fallback về chunk gốc

    tasks = [_retry_chunk(idx, chunk) for idx, chunk in failed_chunks]
    results = asyncio.run(asyncio.gather(*tasks))

    retry_results = {idx: text for idx, text in results}
    still_failed = [
        idx for idx in failed_indices if retry_results.get(idx) == all_chunks[idx]
    ]

    logger.info(
        f"AI Cleanup Retry: {len(failed_indices) - len(still_failed)}/{len(failed_indices)} chunks retry thành công."
    )
    if still_failed:
        logger.warning(
            f"AI Cleanup Retry: {len(still_failed)} chunks vẫn failed sau retry."
        )

    return (retry_results, still_failed)


def _retry_failed_chunks_spell_check(
    failed_indices: List[int],
    all_chunks: List[str],
    api_keys: List[str],
    model_name: str,
    prompt: str,
    ocr_cfg: dict,
) -> tuple[dict[int, str], List[int]]:
    """Retry các chunk failed cho AI Spell Check. Trả về dict {idx: processed_text} và danh sách still_failed."""
    if not failed_indices or not all_chunks:
        return ({}, [])

    spell_check_cfg = ocr_cfg.get("ai_spell_check", {})
    timeout_s = float(spell_check_cfg.get("ai_timeout_seconds", 240))

    failed_chunks = [
        (idx, all_chunks[idx]) for idx in failed_indices if idx < len(all_chunks)
    ]
    logger.info(
        f"AI Spell Check Retry: Đang retry {len(failed_chunks)} chunks failed..."
    )

    # Build safety settings từ config
    safety_level = ocr_cfg.get("safety_level", "BLOCK_ONLY_HIGH")
    safety_settings = _build_safety_settings(safety_level)

    async def _retry_chunk(idx: int, chunk: str) -> tuple[int, str]:
        for key in api_keys:
            try:
                result = await _spell_check_chunk_async(
                    chunk,
                    key,
                    model_name,
                    prompt,
                    idx,
                    len(failed_chunks),
                    timeout_s,
                    safety_settings,
                )
                return (idx, result)
            except Exception:
                continue
        return (idx, chunk)  # Fallback về chunk gốc

    tasks = [_retry_chunk(idx, chunk) for idx, chunk in failed_chunks]
    results = asyncio.run(asyncio.gather(*tasks))

    retry_results = {idx: text for idx, text in results}
    still_failed = [
        idx for idx in failed_indices if retry_results.get(idx) == all_chunks[idx]
    ]

    logger.info(
        f"AI Spell Check Retry: {len(failed_indices) - len(still_failed)}/{len(failed_indices)} chunks retry thành công."
    )
    if still_failed:
        logger.warning(
            f"AI Spell Check Retry: {len(still_failed)} chunks vẫn failed sau retry."
        )

    return (retry_results, still_failed)


def _get_intermediate_file_path(output_path: str, suffix: str) -> str:
    """Tạo đường dẫn file tạm thời dựa trên output_path và suffix."""
    output_dir = os.path.dirname(output_path) if os.path.dirname(output_path) else "."
    output_basename = os.path.basename(output_path)
    output_name_without_ext = os.path.splitext(output_basename)[0]
    return os.path.join(output_dir, output_name_without_ext + suffix)


def _check_existing_files(output_path: str) -> dict:
    """Kiểm tra các file đã tồn tại từ phiên làm việc trước."""
    results = {"ocred": None, "cleanup": None, "output": None, "all_exist": False}

    ocred_path = _get_intermediate_file_path(output_path, "_ocred.txt")
    cleanup_path = _get_intermediate_file_path(output_path, "_cleanup.txt")

    if os.path.exists(ocred_path):
        results["ocred"] = ocred_path
    if os.path.exists(cleanup_path):
        results["cleanup"] = cleanup_path
    if os.path.exists(output_path):
        results["output"] = output_path

    results["all_exist"] = any(
        [results["ocred"], results["cleanup"], results["output"]]
    )
    return results


def _load_resume_file(file_path: str, step_name: str) -> Optional[str]:
    """Load file từ phiên trước để resume."""
    if not os.path.exists(file_path):
        return None
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        logger.info(f"✅ Đã load file {step_name}: {file_path}")
        return content
    except Exception as e:
        logger.warning(f"Không thể load file {step_name} ({file_path}): {e}")
        return None


def _show_completion_menu(
    cleanup_failed: int, spell_check_failed: int, output_path: str = None
) -> str:
    """Hiển thị menu lựa chọn sau khi OCR hoàn tất. Trả về 'retry', 'save', hoặc 'exit'."""
    import threading

    has_failures = cleanup_failed > 0 or spell_check_failed > 0
    user_choice = None
    user_choice_lock = threading.Lock()
    user_choice_done = threading.Event()

    def _auto_save_timer():
        nonlocal user_choice
        time.sleep(600)  # 10 phút = 600 giây
        with user_choice_lock:
            if user_choice is None:
                logger.info("\n⏰ Tự động lưu file sau 10 phút...")
                user_choice = "save"
                user_choice_done.set()

    auto_save_thread = threading.Thread(target=_auto_save_timer, daemon=True)
    auto_save_thread.start()

    if not has_failures:
        # Không có lỗi, chỉ có option save/exit
        logger.info("\n" + "=" * 80)
        logger.info("✅ OCR hoàn tất không có lỗi!")
        logger.info("=" * 80)
        logger.info("Lựa chọn:")
        logger.info("  1. Lưu file (tự động lưu sau 10 phút nếu không chọn)")
        logger.info("  2. Thoát không lưu")
        logger.info("=" * 80)

        while not user_choice_done.is_set():
            try:
                choice = input("\nNhập lựa chọn (1/2): ").strip()
                with user_choice_lock:
                    if choice == "1":
                        user_choice = "save"
                        user_choice_done.set()
                        break
                    elif choice == "2":
                        user_choice = "exit"
                        user_choice_done.set()
                        break
                    else:
                        logger.warning("Lựa chọn không hợp lệ. Vui lòng nhập 1 hoặc 2.")
            except (EOFError, KeyboardInterrupt):
                with user_choice_lock:
                    user_choice = "save"
                    user_choice_done.set()
                break
    else:
        # Có lỗi, hiển thị đầy đủ 3 options
        logger.info("\n" + "=" * 80)
        logger.info("⚠️  OCR hoàn tất với một số lỗi:")
        if cleanup_failed > 0:
            logger.info(f"  - AI Cleanup: {cleanup_failed} chunks failed")
        if spell_check_failed > 0:
            logger.info(f"  - AI Spell Check: {spell_check_failed} chunks failed")
        logger.info("=" * 80)
        logger.info("Lựa chọn:")
        logger.info("  1. Retry các chunk failed")
        logger.info("  2. Lưu file (tự động lưu sau 10 phút nếu không chọn)")
        logger.info("  3. Thoát không lưu")
        logger.info("=" * 80)

        while not user_choice_done.is_set():
            try:
                choice = input("\nNhập lựa chọn (1/2/3): ").strip()
                with user_choice_lock:
                    if choice == "1":
                        user_choice = "retry"
                        user_choice_done.set()
                        break
                    elif choice == "2":
                        user_choice = "save"
                        user_choice_done.set()
                        break
                    elif choice == "3":
                        user_choice = "exit"
                        user_choice_done.set()
                        break
                    else:
                        logger.warning(
                            "Lựa chọn không hợp lệ. Vui lòng nhập 1, 2 hoặc 3."
                        )
            except (EOFError, KeyboardInterrupt):
                with user_choice_lock:
                    user_choice = "save"
                    user_choice_done.set()
                break

    # Đợi user chọn hoặc auto-save
    user_choice_done.wait()
    return user_choice if user_choice else "save"


def ocr_image(image_path: str, config_path: str = "config/config.yaml") -> str:
    ocr_cfg = _detect_bundled_binaries(load_ocr_config(config_path))
    _ensure_dependencies(ocr_cfg)
    if Image is None:
        raise RuntimeError("Pillow not installed. Please install pillow.")
    _apply_tesseract_cfg(ocr_cfg)
    if not os.path.exists(image_path):
        raise FileNotFoundError(image_path)
    logger.info(f"OCR: Đang nhận dạng ảnh: {image_path}")
    img = Image.open(image_path)
    # Auto-detect language/variant nếu cần
    raw_lang = ocr_cfg.get("lang", "vie")
    normalized_lang = _normalize_lang_code(raw_lang)

    # Chỉ detect Chinese variant nếu lang="CN" hoặc "chi" (không có auto-detect)
    needs_chinese_variant_detection = (
        "chi" in normalized_lang.lower()
        and "chi_sim" not in normalized_lang
        and "chi_tra" not in normalized_lang
    )

    if needs_chinese_variant_detection:
        # Chỉ detect Chinese variant (giản thể/phồn thể)
        resolved_lang = _resolve_language(raw_lang, ocr_cfg, sample_img=img)
        text = _image_to_text(img, ocr_cfg, lang_override=resolved_lang)
    else:
        # Chỉ normalize, không cần detect
        resolved_lang = _resolve_language(raw_lang, ocr_cfg, sample_img=None)
        text = _image_to_text(img, ocr_cfg, lang_override=resolved_lang)
    return text


def ocr_pdf(
    pdf_path: str,
    config_path: str = "config/config.yaml",
    pages: Optional[List[int]] = None,
) -> tuple[str, int]:
    ocr_cfg = _detect_bundled_binaries(load_ocr_config(config_path))
    _ensure_dependencies(ocr_cfg)
    if convert_from_path is None:
        raise RuntimeError(
            "pdf2image not installed. Please install pdf2image and poppler if needed."
        )
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(pdf_path)
    _apply_tesseract_cfg(ocr_cfg)

    # Tối ưu DPI: giảm mặc định từ 300 → 250
    dpi = int(ocr_cfg.get("dpi", 250) or 250)
    poppler_path = ocr_cfg.get("poppler_path")

    # Config cho batch processing và memory optimization
    max_batch_size = int(
        ocr_cfg.get("render_batch_size", 20)
    )  # Render tối đa 20 trang/batch
    image_format = ocr_cfg.get("image_format", "jpeg").lower()  # jpeg hoặc png
    jpeg_quality = int(ocr_cfg.get("jpeg_quality", 85))  # Quality 85-90 cho OCR
    memory_optimize = ocr_cfg.get("memory_optimize", True)

    # Resume/caching: sử dụng thư mục cùng tên file input để lưu/trích xuất ảnh các trang
    pdf_p = Path(pdf_path)
    cache_dir = pdf_p.with_suffix("")  # cùng tên với file, bỏ đuôi .pdf

    # Helper function để render và save với batch processing, format tối ưu, và memory management
    def _render_and_save_batch(
        first_page: int,
        last_page: int,
        image_format: str,
        jpeg_quality: int,
        memory_optimize: bool,
    ) -> dict[int, Path]:
        """Render một batch pages và save với format tối ưu. Trả về dict: page_idx → Path."""
        result: dict[int, Path] = {}
        try:
            if poppler_path and isinstance(poppler_path, str) and poppler_path.strip():
                imgs = convert_from_path(
                    pdf_path,
                    dpi=dpi,
                    poppler_path=poppler_path,
                    first_page=first_page,
                    last_page=last_page,
                    thread_count=1,
                )
            else:
                imgs = convert_from_path(
                    pdf_path,
                    dpi=dpi,
                    first_page=first_page,
                    last_page=last_page,
                    thread_count=1,
                )

            for offset, img in enumerate(imgs):
                idx = first_page + offset
                # Chọn extension dựa trên format
                ext = ".jpg" if image_format == "jpeg" else ".png"
                out_path = cache_dir / f"page_{idx:04d}{ext}"

                try:
                    # Save với format và compression tối ưu
                    if image_format == "jpeg":
                        # Convert RGBA/RGB nếu cần (JPEG không hỗ trợ alpha)
                        if img.mode in ("RGBA", "LA", "P"):
                            # Tạo nền trắng cho alpha channel
                            rgb_img = Image.new("RGB", img.size, (255, 255, 255))
                            if img.mode == "P":
                                img = img.convert("RGBA")
                            rgb_img.paste(
                                img,
                                mask=img.split()[-1] if img.mode == "RGBA" else None,
                            )
                            img = rgb_img
                        img.save(
                            str(out_path),
                            format="JPEG",
                            quality=jpeg_quality,
                            optimize=True,
                        )
                    else:
                        # PNG với optimize
                        img.save(str(out_path), format="PNG", optimize=True)

                    result[idx] = out_path

                    # Memory management: giải phóng ngay sau khi save
                    if memory_optimize:
                        del img
                        if offset % 5 == 0:  # Garbage collect mỗi 5 images
                            gc.collect()
                except Exception as e:
                    logger.warning(f"Không thể lưu ảnh cache {out_path}: {e}")

            # Final garbage collect sau batch
            if memory_optimize:
                del imgs
                gc.collect()
        except Exception as e:
            logger.error(f"Render batch {first_page}-{last_page} thất bại: {e}")
        return result

    def _split_range_into_batches(
        range_start: int, range_end: int, batch_size: int
    ) -> List[tuple[int, int]]:
        """Chia một range lớn thành các batches nhỏ."""
        batches = []
        current = range_start
        while current <= range_end:
            batch_end = min(current + batch_size - 1, range_end)
            batches.append((current, batch_end))
            current = batch_end + 1
        return batches

    def _list_cached_images(dir_path: Path) -> List[Path]:
        if not dir_path.exists() or not dir_path.is_dir():
            return []
        image_exts = {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"}
        files = [p for p in dir_path.iterdir() if p.suffix.lower() in image_exts]
        if not files:
            return []

        def sort_key(p: Path):
            name = p.stem
            digits = "".join(ch for ch in name if ch.isdigit())
            return (int(digits) if digits else 0, name)

        return sorted(files, key=sort_key)

    cached_images = _list_cached_images(cache_dir)

    # Lấy tổng số trang PDF để so sánh cache và thực hiện resume nếu thiếu
    def _get_total_pages(pdf_file: str) -> Optional[int]:
        try:
            if pdfplumber is not None:
                with pdfplumber.open(pdf_file) as pdf:
                    return len(pdf.pages)
        except Exception:
            pass
        try:
            if PyPDF2 is not None:
                with open(pdf_file, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    return len(reader.pages)
        except Exception:
            pass
        return None

    total_pages = _get_total_pages(pdf_path)
    if total_pages is not None:
        logger.info(
            f"OCR: PDF có {total_pages} trang. Ảnh cache hiện có: {len(cached_images)}"
        )
    else:
        logger.info(
            f"OCR: Không xác định được tổng số trang. Ảnh cache hiện có: {len(cached_images)}"
        )

    # Filter pages nếu có chỉ định
    if pages and total_pages is not None:
        valid_pages = [p for p in pages if 1 <= p <= total_pages]
        invalid_pages = [p for p in pages if p < 1 or p > total_pages]
        if invalid_pages:
            logger.warning(
                f"Các trang không hợp lệ (nằm ngoài 1-{total_pages}): {invalid_pages}. Bỏ qua."
            )
        if not valid_pages:
            logger.error("Không có trang hợp lệ nào để OCR.")
            return ("", 0)
        pages_to_ocr = sorted(set(valid_pages))
        logger.info(f"OCR: Chỉ OCR {len(pages_to_ocr)} trang: {pages_to_ocr}")
    elif pages:
        # Không biết total_pages nhưng có pages chỉ định → dùng pages đó
        pages_to_ocr = sorted(set([p for p in pages if p > 0]))
        logger.info(
            f"OCR: Chỉ OCR {len(pages_to_ocr)} trang (theo chỉ định): {pages_to_ocr}"
        )
    else:
        pages_to_ocr = None  # Tất cả trang

    # Map chỉ số trang → đường dẫn ảnh (cache) hoặc ảnh render mới
    index_to_image_path: dict[int, Path] = {}
    # Parse chỉ số từ tên ảnh cache kiểu page_0001.png
    for p in cached_images:
        name = p.stem
        digits = "".join(ch for ch in name if ch.isdigit())
        if digits:
            try:
                idx = int(digits)
                # Chỉ lấy cache nếu trang đó nằm trong pages_to_ocr (hoặc pages_to_ocr = None)
                if pages_to_ocr is None or idx in pages_to_ocr:
                    index_to_image_path[idx] = p
            except Exception:
                continue

    # Render bổ sung cho các trang thiếu nếu biết total_pages
    if total_pages is not None:
        # Tính missing pages: nếu có pages_to_ocr thì chỉ tính trong đó, ngược lại tính tất cả
        if pages_to_ocr is not None:
            target_pages = set(pages_to_ocr)
            missing = [i for i in target_pages if i not in index_to_image_path]
        else:
            missing = [
                i for i in range(1, total_pages + 1) if i not in index_to_image_path
            ]
        if missing:
            logger.info(
                f"OCR: Phát hiện thiếu {len(missing)}/{total_pages} ảnh → render phần còn thiếu"
            )
            # Tạo thư mục cache nếu cần
            try:
                cache_dir.mkdir(parents=True, exist_ok=True)
            except Exception:
                pass
            # Gom missing pages thành các khoảng liên tiếp để render theo range
            ranges: List[tuple[int, int]] = []
            start = prev = None
            for m in missing:
                if start is None:
                    start = prev = m
                elif m == prev + 1:
                    prev = m
                else:
                    ranges.append((start, prev))
                    start = prev = m
            if start is not None:
                ranges.append((start, prev))

            # Render với batch processing để giảm memory usage
            for first, last in ranges:
                # Chia range lớn thành batches nhỏ
                batches = _split_range_into_batches(first, last, max_batch_size)
                for batch_first, batch_last in batches:
                    logger.info(
                        f"OCR: Render bổ sung trang {batch_first}–{batch_last}/{last} (dpi={dpi}, format={image_format})"
                    )
                    batch_results = _render_and_save_batch(
                        batch_first,
                        batch_last,
                        image_format,
                        jpeg_quality,
                        memory_optimize,
                    )
                    index_to_image_path.update(batch_results)

    # Nếu vẫn chưa có ảnh nào (không có cache và không biết total), render
    if not index_to_image_path:
        if pages_to_ocr is not None and total_pages is not None:
            # Chỉ render các trang được chỉ định
            logger.info(
                f"OCR: Chuyển PDF → ảnh (dpi={dpi}) cho {len(pages_to_ocr)} trang: {pages_to_ocr}"
            )
            # Gom pages_to_ocr thành ranges để render hiệu quả
            ranges: List[tuple[int, int]] = []
            pages_sorted = sorted(pages_to_ocr)
            start = prev = None
            for p in pages_sorted:
                if start is None:
                    start = prev = p
                elif p == prev + 1:
                    prev = p
                else:
                    ranges.append((start, prev))
                    start = prev = p
            if start is not None:
                ranges.append((start, prev))

            try:
                cache_dir.mkdir(parents=True, exist_ok=True)
            except Exception:
                pass

            # Render với batch processing để giảm memory usage
            for first, last in ranges:
                # Chia range lớn thành batches nhỏ
                batches = _split_range_into_batches(first, last, max_batch_size)
                for batch_first, batch_last in batches:
                    logger.info(
                        f"OCR: Render trang {batch_first}–{batch_last}/{last} (dpi={dpi}, format={image_format})"
                    )
                    batch_results = _render_and_save_batch(
                        batch_first,
                        batch_last,
                        image_format,
                        jpeg_quality,
                        memory_optimize,
                    )
                    index_to_image_path.update(batch_results)
        else:
            # Render toàn bộ (không có pages filter) - CHIA THÀNH BATCHES để tránh ngốn RAM
            if total_pages is not None:
                logger.info(
                    f"OCR: Chuyển PDF → ảnh (dpi={dpi}, format={image_format}): {total_pages} trang - render theo batch {max_batch_size} trang/batch"
                )
                try:
                    cache_dir.mkdir(parents=True, exist_ok=True)
                except Exception:
                    pass
                # Chia toàn bộ PDF thành batches
                batches = _split_range_into_batches(1, total_pages, max_batch_size)
                for batch_first, batch_last in batches:
                    logger.info(
                        f"OCR: Render batch {batch_first}–{batch_last}/{total_pages}"
                    )
                    batch_results = _render_and_save_batch(
                        batch_first,
                        batch_last,
                        image_format,
                        jpeg_quality,
                        memory_optimize,
                    )
                    index_to_image_path.update(batch_results)
            else:
                # Không biết total_pages, phải render hết (vẫn cố gắng dùng thread_count=1 để giảm memory)
                logger.info(
                    f"OCR: Chuyển PDF → ảnh (dpi={dpi}, format={image_format}): không biết số trang, render toàn bộ"
                )
                try:
                    cache_dir.mkdir(parents=True, exist_ok=True)
                except Exception:
                    pass
                if (
                    poppler_path
                    and isinstance(poppler_path, str)
                    and poppler_path.strip()
                ):
                    all_imgs = convert_from_path(
                        pdf_path, dpi=dpi, poppler_path=poppler_path, thread_count=1
                    )
                else:
                    all_imgs = convert_from_path(pdf_path, dpi=dpi, thread_count=1)

                # Save với format tối ưu và memory management
                ext = ".jpg" if image_format == "jpeg" else ".png"
                for idx, img in enumerate(all_imgs, start=1):
                    out_path = cache_dir / f"page_{idx:04d}{ext}"
                    try:
                        if image_format == "jpeg":
                            if img.mode in ("RGBA", "LA", "P"):
                                rgb_img = Image.new("RGB", img.size, (255, 255, 255))
                                if img.mode == "P":
                                    img = img.convert("RGBA")
                                rgb_img.paste(
                                    img,
                                    mask=img.split()[-1]
                                    if img.mode == "RGBA"
                                    else None,
                                )
                                img = rgb_img
                            img.save(
                                str(out_path),
                                format="JPEG",
                                quality=jpeg_quality,
                                optimize=True,
                            )
                        else:
                            img.save(str(out_path), format="PNG", optimize=True)
                        index_to_image_path[idx] = out_path
                        if memory_optimize:
                            del img
                            if idx % 5 == 0:
                                gc.collect()
                    except Exception as e:
                        logger.warning(f"Không thể lưu ảnh cache {out_path}: {e}")
                if memory_optimize:
                    del all_imgs
                    gc.collect()

    # Tạo danh sách ảnh theo thứ tự trang để OCR
    if pages_to_ocr is not None:
        # Chỉ OCR các trang được chỉ định (và có sẵn ảnh)
        ordered_indices = sorted(
            [idx for idx in pages_to_ocr if idx in index_to_image_path]
        )
    elif total_pages is None:
        # Không biết tổng trang: dùng thứ tự theo index hiện có
        ordered_indices = sorted(index_to_image_path.keys())
    else:
        ordered_indices = list(range(1, total_pages + 1))

    # Auto-detect language/variant nếu cần (chỉ một lần cho toàn bộ PDF)
    resolved_lang = None
    raw_lang = ocr_cfg.get("lang", "vie")
    # Normalize trước để check "auto" và "chi"
    normalized_lang = _normalize_lang_code(raw_lang)

    # Chỉ detect Chinese variant nếu lang="CN" hoặc "chi" (không có auto-detect)
    needs_chinese_variant_detection = (
        "chi" in normalized_lang.lower()
        and "chi_sim" not in normalized_lang
        and "chi_tra" not in normalized_lang
    )

    if needs_chinese_variant_detection and ordered_indices:
        # Chinese variant detection: chỉ cần 1 trang đầu để detect giản thể/phồn thể
        first_page_idx = ordered_indices[0]
        first_page_path = index_to_image_path.get(first_page_idx)
        if first_page_path:
            try:
                logger.info(
                    f"Đang nhận biết Chinese variant (giản thể/phồn thể) từ trang đầu (lang config: {raw_lang})..."
                )
                with Image.open(str(first_page_path)) as sample_img:
                    resolved_lang = _resolve_language(
                        raw_lang, ocr_cfg, sample_img=sample_img
                    )
                logger.info(f"Đã detect Chinese variant: {resolved_lang}")
            except Exception as e:
                logger.warning(
                    f"Không thể detect Chinese variant từ trang đầu: {e}. Dùng mặc định chi_sim."
                )
                resolved_lang = _resolve_language(raw_lang, ocr_cfg, sample_img=None)
        else:
            resolved_lang = _resolve_language(raw_lang, ocr_cfg, sample_img=None)
    else:
        # Không cần detect variant, chỉ normalize
        resolved_lang = _resolve_language(raw_lang, ocr_cfg, sample_img=None)

    texts: List[str] = []
    total = len(ordered_indices)
    logger.info(f"OCR: Tổng số trang cần xử lý: {total}")
    show_progress = bool(ocr_cfg.get("show_progress", True))
    progress_interval = float(ocr_cfg.get("progress_log_interval_seconds", 60))
    if show_progress and tqdm is not None and total > 1:
        start_ts = time.time()
        with tqdm(total=total, desc="OCR PDF", unit="trang") as pbar:
            for i, page_idx in enumerate(ordered_indices, start=1):
                p = index_to_image_path.get(page_idx)
                if p is None:
                    logger.warning(f"Thiếu ảnh cho trang {page_idx}, bỏ qua")
                    pbar.update(1)
                    continue
                try:
                    # Thử mở ảnh với LOAD_TRUNCATED_IMAGES để xử lý ảnh bị truncated
                    try:
                        with Image.open(str(p)) as img:
                            # Thử load full image với LOAD_TRUNCATED_IMAGES nếu bị truncated
                            img.load()  # Load toàn bộ image data
                            text = _image_to_text(
                                img, ocr_cfg, lang_override=resolved_lang
                            )
                    except Exception as load_error:
                        # Nếu vẫn lỗi, thử với LOAD_TRUNCATED_IMAGES flag
                        if "truncated" in str(load_error).lower():
                            logger.warning(
                                f"Ảnh trang {page_idx} bị truncated, thử load với LOAD_TRUNCATED_IMAGES..."
                            )
                            with Image.open(str(p)) as img:
                                # Pillow tự động xử lý truncated images nếu có thể
                                try:
                                    # Thử verify=False để bỏ qua một số checks
                                    img.verify()
                                    img = Image.open(str(p))  # Reopen sau verify
                                    text = _image_to_text(
                                        img, ocr_cfg, lang_override=resolved_lang
                                    )
                                except Exception:
                                    # Nếu vẫn lỗi, thử render lại từ PDF
                                    logger.warning(
                                        f"Không thể load ảnh truncated trang {page_idx}, thử render lại từ PDF..."
                                    )
                                    try:
                                        # Xóa file cache bị lỗi
                                        if p.exists():
                                            try:
                                                p.unlink()
                                                logger.debug(
                                                    f"Đã xóa file cache bị lỗi: {p}"
                                                )
                                            except Exception:
                                                pass

                                        # Render lại từ PDF (single page - dùng format tối ưu)
                                        if (
                                            poppler_path
                                            and isinstance(poppler_path, str)
                                            and poppler_path.strip()
                                        ):
                                            imgs = convert_from_path(
                                                pdf_path,
                                                dpi=dpi,
                                                poppler_path=poppler_path,
                                                first_page=page_idx,
                                                last_page=page_idx,
                                                thread_count=1,
                                            )
                                        else:
                                            imgs = convert_from_path(
                                                pdf_path,
                                                dpi=dpi,
                                                first_page=page_idx,
                                                last_page=page_idx,
                                                thread_count=1,
                                            )

                                        if imgs and len(imgs) > 0:
                                            img = imgs[0]
                                            # Lưu lại vào cache với format tối ưu
                                            try:
                                                cache_dir.mkdir(
                                                    parents=True, exist_ok=True
                                                )
                                                if image_format == "jpeg":
                                                    if img.mode in ("RGBA", "LA", "P"):
                                                        rgb_img = Image.new(
                                                            "RGB",
                                                            img.size,
                                                            (255, 255, 255),
                                                        )
                                                        if img.mode == "P":
                                                            img = img.convert("RGBA")
                                                        rgb_img.paste(
                                                            img,
                                                            mask=img.split()[-1]
                                                            if img.mode == "RGBA"
                                                            else None,
                                                        )
                                                        img = rgb_img
                                                    img.save(
                                                        str(p),
                                                        format="JPEG",
                                                        quality=jpeg_quality,
                                                        optimize=True,
                                                    )
                                                else:
                                                    img.save(
                                                        str(p),
                                                        format="PNG",
                                                        optimize=True,
                                                    )
                                                logger.info(
                                                    f"Đã render lại và lưu cache cho trang {page_idx}"
                                                )
                                            except Exception as save_err:
                                                logger.debug(
                                                    f"Không thể lưu cache lại: {save_err}"
                                                )

                                            # OCR lại (trước khi giải phóng memory)
                                            text = _image_to_text(
                                                img,
                                                ocr_cfg,
                                                lang_override=resolved_lang,
                                            )

                                            # Giải phóng memory sau khi OCR xong
                                            if memory_optimize:
                                                del img, imgs
                                                gc.collect()
                                        else:
                                            logger.warning(
                                                f"Không thể render lại trang {page_idx} từ PDF"
                                            )
                                            text = ""
                                    except Exception as render_error:
                                        logger.warning(
                                            f"Không thể render lại trang {page_idx}: {render_error}"
                                        )
                                        text = ""
                        else:
                            raise load_error  # Nếu không phải truncated error, re-raise
                except Exception as e:
                    logger.warning(
                        f"Không thể mở/OCR ảnh cho trang {page_idx} ({p}): {e}"
                    )
                    text = ""
                texts.append(text)
                elapsed = time.time() - start_ts
                avg = elapsed / i if i > 0 else 0.0
                remaining = max(total - i, 0) * avg
                pbar.set_postfix(avg_s_per_page=f"{avg:.2f}", eta=f"{remaining:.0f}s")
                pbar.update(1)
    else:
        start_ts = time.time()
        last_log = start_ts
        for i, page_idx in enumerate(ordered_indices, start=1):
            p = index_to_image_path.get(page_idx)
            if p is None:
                logger.warning(f"Thiếu ảnh cho trang {page_idx}, bỏ qua")
                continue
            try:
                # Thử mở ảnh với xử lý truncated images
                try:
                    with Image.open(str(p)) as img:
                        img.load()  # Load toàn bộ image data
                        texts.append(
                            _image_to_text(img, ocr_cfg, lang_override=resolved_lang)
                        )
                except Exception as load_error:
                    # Nếu bị truncated, thử các cách khắc phục
                    if "truncated" in str(load_error).lower():
                        logger.warning(
                            f"Ảnh trang {page_idx} bị truncated, thử load với LOAD_TRUNCATED_IMAGES..."
                        )
                        try:
                            with Image.open(str(p)) as img:
                                img.verify()
                                img = Image.open(str(p))  # Reopen sau verify
                                texts.append(
                                    _image_to_text(
                                        img, ocr_cfg, lang_override=resolved_lang
                                    )
                                )
                        except Exception:
                            # Thử render lại từ PDF
                            logger.warning(
                                f"Không thể load ảnh truncated trang {page_idx}, thử render lại từ PDF..."
                            )
                            try:
                                # Xóa file cache bị lỗi
                                if p.exists():
                                    try:
                                        p.unlink()
                                    except Exception:
                                        pass

                                # Render lại từ PDF (single page - dùng format tối ưu)
                                if (
                                    poppler_path
                                    and isinstance(poppler_path, str)
                                    and poppler_path.strip()
                                ):
                                    imgs = convert_from_path(
                                        pdf_path,
                                        dpi=dpi,
                                        poppler_path=poppler_path,
                                        first_page=page_idx,
                                        last_page=page_idx,
                                        thread_count=1,
                                    )
                                else:
                                    imgs = convert_from_path(
                                        pdf_path,
                                        dpi=dpi,
                                        first_page=page_idx,
                                        last_page=page_idx,
                                        thread_count=1,
                                    )

                                if imgs and len(imgs) > 0:
                                    img = imgs[0]
                                    # Lưu lại vào cache với format tối ưu
                                    try:
                                        cache_dir.mkdir(parents=True, exist_ok=True)
                                        if image_format == "jpeg":
                                            if img.mode in ("RGBA", "LA", "P"):
                                                rgb_img = Image.new(
                                                    "RGB", img.size, (255, 255, 255)
                                                )
                                                if img.mode == "P":
                                                    img = img.convert("RGBA")
                                                rgb_img.paste(
                                                    img,
                                                    mask=img.split()[-1]
                                                    if img.mode == "RGBA"
                                                    else None,
                                                )
                                                img = rgb_img
                                            img.save(
                                                str(p),
                                                format="JPEG",
                                                quality=jpeg_quality,
                                                optimize=True,
                                            )
                                        else:
                                            img.save(
                                                str(p), format="PNG", optimize=True
                                            )
                                    except Exception:
                                        pass

                                    # OCR lại (trước khi giải phóng memory)
                                    texts.append(
                                        _image_to_text(
                                            img, ocr_cfg, lang_override=resolved_lang
                                        )
                                    )

                                    # Giải phóng memory sau khi OCR xong
                                    if memory_optimize:
                                        del img, imgs
                                        gc.collect()
                                else:
                                    logger.warning(
                                        f"Không thể render lại trang {page_idx} từ PDF"
                                    )
                                    texts.append("")
                            except Exception as render_error:
                                logger.warning(
                                    f"Không thể render lại trang {page_idx}: {render_error}"
                                )
                                texts.append("")
                    else:
                        raise load_error
            except Exception as e:
                logger.warning(f"Không thể mở/OCR ảnh cho trang {page_idx} ({p}): {e}")
                texts.append("")
            now = time.time()
            if now - last_log >= max(5.0, progress_interval):  # báo cáo định kỳ
                elapsed = now - start_ts
                avg = elapsed / i if i > 0 else 0.0
                remaining = max(total - i, 0) * avg
                logger.info(
                    f"OCR: Trang {i}/{total} • TB {avg:.2f}s/trang • ETA ~{remaining:.0f}s"
                )
                last_log = now
        elapsed = time.time() - start_ts
        avg = elapsed / total if total > 0 else 0.0
        logger.info(f"OCR: Hoàn tất {total} trang • TB {avg:.2f}s/trang")
    # Trả về số trang đã thực sự OCR (không phải total_pages)
    pages_processed = len(ordered_indices)
    return ("\n\n".join(texts), pages_processed)


def ocr_file(
    input_path: str,
    config_path: str = "config/config.yaml",
    pages: Optional[List[int]] = None,
    output_path: Optional[str] = None,
    skip_steps: Optional[dict] = None,
) -> str:
    """
    Extract text từ file PDF hoặc ảnh.
    Tự động detect PDF scan vs text-based để tối ưu.

    Args:
        output_path: Đường dẫn file output (để tạo tên file tạm thời)
        skip_steps: Dict với keys 'ocr', 'cleanup', 'spell_check' để skip các bước đã hoàn tất
    """
    _ensure_logger_config()
    pipeline_start_time = time.time()
    total_pages_processed = 0

    if skip_steps is None:
        skip_steps = {}

    if not os.path.exists(input_path):
        raise FileNotFoundError(input_path)

    ocr_cfg = _detect_bundled_binaries(load_ocr_config(config_path))
    _ensure_dependencies(ocr_cfg)

    ext = os.path.splitext(input_path)[-1].lower()

    # Xử lý PDF
    if ext == ".pdf":
        auto_detect = ocr_cfg.get("auto_detect_pdf_type", True)

        if auto_detect:
            logger.info(f"Đang phát hiện loại PDF: {input_path}")
            pdf_type = detect_pdf_type(input_path)
            logger.info(f"PDF type: {pdf_type}")

            if pdf_type == "text":
                logger.info("PDF có text layer → Extract text trực tiếp (nhanh)")
                text = extract_text_from_pdf(input_path, ocr_cfg, pages)
                # Đếm số trang đã xử lý
                if pages:
                    # Validate và đếm số trang hợp lệ
                    try:
                        if pdfplumber is not None:
                            with pdfplumber.open(input_path) as pdf:
                                total = len(pdf.pages)
                                valid_pages = [p for p in pages if 1 <= p <= total]
                                total_pages_processed = len(valid_pages)
                        elif PyPDF2 is not None:
                            with open(input_path, "rb") as f:
                                reader = PyPDF2.PdfReader(f)
                                total = len(reader.pages)
                                valid_pages = [p for p in pages if 1 <= p <= total]
                                total_pages_processed = len(valid_pages)
                        else:
                            total_pages_processed = len(pages)  # Fallback
                    except Exception:
                        total_pages_processed = len(pages)  # Fallback
                else:
                    try:
                        if pdfplumber is not None:
                            with pdfplumber.open(input_path) as pdf:
                                total_pages_processed = len(pdf.pages)
                        elif PyPDF2 is not None:
                            with open(input_path, "rb") as f:
                                reader = PyPDF2.PdfReader(f)
                                total_pages_processed = len(reader.pages)
                    except Exception:
                        total_pages_processed = 0
            else:
                logger.info("PDF scan → Sử dụng OCR")
                text, total_pages_processed = ocr_pdf(input_path, config_path, pages)
        else:
            # Force OCR nếu auto_detect = false
            text, total_pages_processed = ocr_pdf(input_path, config_path, pages)
    # Xử lý ảnh
    elif ext in {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff"}:
        text = ocr_image(input_path, config_path)
        total_pages_processed = 1  # Một ảnh = 1 trang
    else:
        raise ValueError(f"Unsupported input format for OCR: {ext}")

    # Lưu file sau bước OCR nếu chưa skip và có output_path
    if not skip_steps.get("ocr", False) and output_path:
        ocred_path = _get_intermediate_file_path(output_path, "_ocred.txt")
        try:
            with open(ocred_path, "w", encoding="utf-8") as f:
                f.write(text)
            logger.info(f"💾 Đã lưu kết quả OCR: {ocred_path}")
        except Exception as e:
            logger.warning(f"Không thể lưu file OCR: {e}")

    # Áp dụng AI cleanup nếu enabled
    cleanup_cfg = ocr_cfg.get("ai_cleanup", {})
    cleanup_failed = 0
    cleanup_failed_indices = []
    cleanup_original_chunks = []

    if cleanup_cfg.get("enabled", False) and not skip_steps.get("cleanup", False):
        result = ai_cleanup_text(text, ocr_cfg)
        if isinstance(result, tuple):
            text, cleanup_failed_indices, cleanup_original_chunks = result
            cleanup_failed = len(cleanup_failed_indices)
        else:
            text = result

        # Lưu file sau bước cleanup nếu có output_path
        if output_path:
            cleanup_path = _get_intermediate_file_path(output_path, "_cleanup.txt")
            try:
                with open(cleanup_path, "w", encoding="utf-8") as f:
                    f.write(text)
                logger.info(f"💾 Đã lưu kết quả Cleanup: {cleanup_path}")
            except Exception as e:
                logger.warning(f"Không thể lưu file Cleanup: {e}")
    elif skip_steps.get("cleanup", False):
        logger.info("⏭️  Bỏ qua bước Cleanup (đã có file từ phiên trước)")

    # Áp dụng AI spell check và paragraph restoration nếu enabled
    spell_check_cfg = ocr_cfg.get("ai_spell_check", {})
    spell_check_failed = 0
    spell_check_failed_indices = []
    spell_check_original_chunks = []

    if spell_check_cfg.get("enabled", False) and not skip_steps.get(
        "spell_check", False
    ):
        result = ai_spell_check_and_paragraph_restore(text, ocr_cfg)
        if isinstance(result, tuple):
            text, spell_check_failed_indices, spell_check_original_chunks = result
            spell_check_failed = len(spell_check_failed_indices)
        else:
            text = result
    elif skip_steps.get("spell_check", False):
        logger.info("⏭️  Bỏ qua bước Spell Check (đã có file từ phiên trước)")

    # Log tổng kết
    total_time = time.time() - pipeline_start_time
    hours = int(total_time // 3600)
    minutes = int((total_time % 3600) // 60)
    seconds = int(total_time % 60)
    time_str = (
        f"{hours:02d}:{minutes:02d}:{seconds:02d}"
        if hours > 0
        else f"{minutes:02d}:{seconds:02d}"
    )

    logger.info("=" * 80)
    logger.info("📊 TỔNG KẾT OCR PIPELINE")
    logger.info(f"⏱️  Tổng thời gian: {time_str} ({total_time:.2f} giây)")
    logger.info(f"📄 Số trang đã OCR: {total_pages_processed}")
    if cleanup_cfg.get("enabled", False):
        if cleanup_failed > 0:
            logger.info(
                f"🧹 AI Cleanup: {cleanup_failed} chunks failed (đã lưu nội dung gốc)"
            )
        else:
            logger.info("🧹 AI Cleanup: Hoàn tất không có lỗi")
    if spell_check_cfg.get("enabled", False):
        if spell_check_failed > 0:
            logger.info(
                f"✅ AI Spell Check: {spell_check_failed} chunks failed (đã lưu nội dung gốc)"
            )
        else:
            logger.info("✅ AI Spell Check: Hoàn tất không có lỗi")
    logger.info("=" * 80)

    # Lưu lại cấu trúc chunks để có thể merge lại sau retry
    cleanup_all_chunks = cleanup_original_chunks if cleanup_original_chunks else []
    spell_check_all_chunks = (
        spell_check_original_chunks if spell_check_original_chunks else []
    )

    # Trả về text và thông tin failures để menu xử lý
    return {
        "text": text,
        "cleanup_failed": cleanup_failed,
        "cleanup_failed_indices": cleanup_failed_indices,
        "cleanup_original_chunks": cleanup_original_chunks,
        "cleanup_all_chunks": cleanup_all_chunks,  # Tất cả chunks (để merge lại)
        "spell_check_failed": spell_check_failed,
        "spell_check_failed_indices": spell_check_failed_indices,
        "spell_check_original_chunks": spell_check_original_chunks,
        "spell_check_all_chunks": spell_check_all_chunks,  # Tất cả chunks (để merge lại)
        "ocr_cfg": ocr_cfg,
    }


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="OCR scan file (PDF/image) and extract text"
    )
    parser.add_argument("input", help="Path to image or PDF file (scan)")
    parser.add_argument(
        "--config", default="config/config.yaml", help="Path to YAML config"
    )
    parser.add_argument("--output", help="Save recognized text to file")
    parser.add_argument(
        "--pages",
        help="Chỉ định các trang cần OCR. Ví dụ: '1,2,5,7' hoặc '1-7' hoặc '1-3,5,7-9'",
    )
    parser.add_argument(
        "--format",
        choices=["txt", "docx"],
        default=None,
        help="Định dạng file output (txt hoặc docx). Nếu không chỉ định, sẽ tự động detect từ extension của --output",
    )
    args = parser.parse_args()

    _ensure_logger_config()

    # Xác định output_path và output_format
    if args.output:
        output_path = args.output
        # Detect format từ extension nếu --format không được chỉ định
        if args.format:
            output_format = args.format
        else:
            ext = os.path.splitext(output_path)[1].lower()
            if ext == ".docx":
                output_format = "docx"
            else:
                output_format = "txt"
    else:
        # Tạo output_path mặc định từ input_path
        input_dir = os.path.dirname(args.input) if os.path.dirname(args.input) else "."
        input_basename = os.path.basename(args.input)
        input_name_without_ext = os.path.splitext(input_basename)[0]
        if args.format == "docx":
            output_path = os.path.join(
                input_dir, input_name_without_ext + "_ocr_result.docx"
            )
            output_format = "docx"
        else:
            output_path = os.path.join(
                input_dir, input_name_without_ext + "_ocr_result.txt"
            )
            output_format = "txt"

    # Check & Resume: Kiểm tra file từ phiên trước
    existing_files = _check_existing_files(output_path)
    skip_steps = {}

    if existing_files["all_exist"]:
        logger.info("\n" + "=" * 80)
        logger.info("🔍 PHÁT HIỆN FILE TỪ PHIÊN LÀM VIỆC TRƯỚC")
        logger.info("=" * 80)

        found_files = []
        if existing_files["ocred"]:
            found_files.append(f"  • File OCR: {existing_files['ocred']}")
        if existing_files["cleanup"]:
            found_files.append(f"  • File Cleanup: {existing_files['cleanup']}")
        if existing_files["output"]:
            found_files.append(f"  • File Output: {existing_files['output']}")

        if found_files:
            logger.info("\nCác file đã phát hiện:")
            for f in found_files:
                logger.info(f)
            logger.info("")

            logger.info("Bạn có muốn sử dụng các file này để tiếp tục?")
            logger.info("")
            logger.info("  1. Có, tiếp tục từ Cleanup")
            logger.info("  2. Có, tiếp tục từ Spell Check")
            logger.info("  3. Có, nhưng chạy lại toàn bộ (OCR + Cleanup + Spell Check)")
            logger.info("  4. Không, kết thúc tác vụ")
            logger.info("")

            while True:
                try:
                    choice = input("Nhập lựa chọn (1/2/3/4): ").strip()
                    if choice == "1":
                        # Resume từ Cleanup (cần file _ocred.txt)
                        if existing_files["ocred"]:
                            skip_steps["ocr"] = True
                            logger.info("⏭️  Sẽ bỏ qua OCR, tiếp tục từ Cleanup")
                        else:
                            logger.warning(
                                "⚠️  Không tìm thấy file OCR (_ocred.txt). Không thể tiếp tục từ Cleanup."
                            )
                            logger.info("🔄 Sẽ chạy lại toàn bộ quy trình")
                            skip_steps = {}
                        break
                    elif choice == "2":
                        # Resume từ Spell Check (cần file _cleanup.txt)
                        if existing_files["cleanup"]:
                            skip_steps["ocr"] = True
                            skip_steps["cleanup"] = True
                            logger.info(
                                "⏭️  Sẽ bỏ qua OCR và Cleanup, chỉ chạy Spell Check"
                            )
                        else:
                            logger.warning(
                                "⚠️  Không tìm thấy file Cleanup (_cleanup.txt). Không thể tiếp tục từ Spell Check."
                            )
                            if existing_files["ocred"]:
                                logger.info(
                                    "💡 Phát hiện file OCR. Sẽ tiếp tục từ Cleanup."
                                )
                                skip_steps["ocr"] = True
                            else:
                                logger.info("🔄 Sẽ chạy lại toàn bộ quy trình")
                                skip_steps = {}
                        break
                    elif choice == "3":
                        # Chạy lại toàn bộ
                        skip_steps = {}
                        logger.info("🔄 Sẽ chạy lại toàn bộ quy trình")
                        break
                    elif choice == "4":
                        logger.info("Kết thúc tác vụ.")
                        sys.exit(0)
                    else:
                        logger.warning("Vui lòng nhập 1, 2, 3 hoặc 4.")
                except (KeyboardInterrupt, EOFError):
                    logger.info("\nKết thúc tác vụ.")
                    sys.exit(0)

    logger.info("Bắt đầu OCR pipeline...")

    # Parse pages nếu có
    pages_list = None
    if args.pages:
        pages_list = _parse_pages(args.pages)
        if pages_list:
            logger.info(f"Chỉ OCR các trang: {pages_list}")
        else:
            logger.warning(
                f"Không parse được pages từ '{args.pages}'. Sẽ OCR tất cả trang."
            )

    # Load file từ phiên trước nếu resume
    initial_text = None
    if skip_steps.get("ocr", False):
        if skip_steps.get("cleanup", False):
            # Resume từ Spell Check → load file Cleanup
            if existing_files["cleanup"]:
                initial_text = _load_resume_file(existing_files["cleanup"], "Cleanup")
            else:
                logger.error("❌ Không tìm thấy file Cleanup để resume từ Spell Check!")
                logger.info("🔄 Sẽ chạy lại toàn bộ quy trình")
                skip_steps = {}
        else:
            # Resume từ Cleanup → load file OCR
            if existing_files["ocred"]:
                initial_text = _load_resume_file(existing_files["ocred"], "OCR")
            else:
                logger.error("❌ Không tìm thấy file OCR để resume từ Cleanup!")
                logger.info("🔄 Sẽ chạy lại toàn bộ quy trình")
                skip_steps = {}

    # Chạy pipeline với skip_steps
    if initial_text:
        # Resume từ file đã có → chỉ cần chạy các bước còn lại
        ocr_cfg = _detect_bundled_binaries(load_ocr_config(args.config))
        _ensure_dependencies(ocr_cfg)
        text = initial_text

        cleanup_cfg = ocr_cfg.get("ai_cleanup", {})
        cleanup_failed = 0
        cleanup_failed_indices = []
        cleanup_original_chunks = []

        # Chạy cleanup nếu cần (không skip)
        if cleanup_cfg.get("enabled", False) and not skip_steps.get("cleanup", False):
            result = ai_cleanup_text(text, ocr_cfg)
            if isinstance(result, tuple):
                text, cleanup_failed_indices, cleanup_original_chunks = result
                cleanup_failed = len(cleanup_failed_indices)
            else:
                text = result

            # Lưu file sau cleanup
            cleanup_path = _get_intermediate_file_path(output_path, "_cleanup.txt")
            try:
                with open(cleanup_path, "w", encoding="utf-8") as f:
                    f.write(text)
                logger.info(f"💾 Đã lưu kết quả Cleanup: {cleanup_path}")
            except Exception as e:
                logger.warning(f"Không thể lưu file Cleanup: {e}")

        # Chạy spell check nếu cần
        spell_check_cfg = ocr_cfg.get("ai_spell_check", {})
        spell_check_failed = 0
        spell_check_failed_indices = []
        spell_check_original_chunks = []

        if spell_check_cfg.get("enabled", False):
            result = ai_spell_check_and_paragraph_restore(text, ocr_cfg)
            if isinstance(result, tuple):
                text, spell_check_failed_indices, spell_check_original_chunks = result
                spell_check_failed = len(spell_check_failed_indices)
            else:
                text = result

        # Tạo result dict giống format của ocr_file
        result = {
            "text": text,
            "cleanup_failed": cleanup_failed,
            "cleanup_failed_indices": cleanup_failed_indices,
            "cleanup_original_chunks": cleanup_original_chunks,
            "cleanup_all_chunks": cleanup_original_chunks
            if cleanup_original_chunks
            else [],
            "spell_check_failed": spell_check_failed,
            "spell_check_failed_indices": spell_check_failed_indices,
            "spell_check_original_chunks": spell_check_original_chunks,
            "spell_check_all_chunks": spell_check_original_chunks
            if spell_check_original_chunks
            else [],
            "ocr_cfg": ocr_cfg,
        }
    else:
        # Kiểm tra nếu cần xuất DOCX với images
        if (
            output_format == "docx"
            and os.path.splitext(args.input)[1].lower() == ".pdf"
        ):
            # Chỉ hỗ trợ DOCX cho PDF có text layer
            ocr_cfg = _detect_bundled_binaries(load_ocr_config(args.config))
            _ensure_dependencies(ocr_cfg)

            pdf_type = detect_pdf_type(args.input)

            if pdf_type == "text":
                # PDF có text layer → extract text và images trực tiếp
                logger.info("📄 PDF có text layer → Tạo DOCX với text và images")
                try:
                    create_docx_from_pdf(args.input, output_path, ocr_cfg, pages_list)
                    logger.info(f"✅ Đã tạo DOCX thành công: {output_path}")
                    logger.info("Hoàn tất OCR pipeline.")
                    sys.exit(0)
                except Exception as e:
                    logger.error(f"❌ Lỗi khi tạo DOCX: {e}")
                    logger.warning("Fallback về định dạng TXT...")
                    output_format = "txt"
                    # Cập nhật output_path để có extension .txt
                    output_path = os.path.splitext(output_path)[0] + ".txt"
            else:
                # PDF scan → không thể extract images, chỉ có thể OCR text
                logger.warning(
                    "⚠️  PDF scan không thể extract images. Chỉ có thể xuất text."
                )
                logger.warning(
                    "💡 Để xuất DOCX với images, cần PDF có text layer (text-based PDF)."
                )
                logger.warning("🔄 Fallback về định dạng TXT...")
                output_format = "txt"
                # Cập nhật output_path để có extension .txt
                output_path = os.path.splitext(output_path)[0] + ".txt"

        # Chạy toàn bộ pipeline (normal OCR → TXT)
        result = ocr_file(
            args.input,
            config_path=args.config,
            pages=pages_list,
            output_path=output_path,
            skip_steps=skip_steps,
        )

    # Xử lý kết quả từ ocr_file
    result_text = result["text"]
    cleanup_failed = result["cleanup_failed"]
    spell_check_failed = result["spell_check_failed"]
    ocr_cfg = result["ocr_cfg"]

    # Hiển thị menu completion
    if output_path:
        logger.info(f"\n📁 File output: {output_path}")
        user_choice = _show_completion_menu(
            cleanup_failed, spell_check_failed, output_path
        )

        if user_choice == "retry":
            # Retry các chunk failed và merge lại (logic giống phần else)
            cleanup_cfg = ocr_cfg.get("ai_cleanup", {})
            spell_check_cfg = ocr_cfg.get("ai_spell_check", {})
            updated_text = result_text

            if (
                result["cleanup_failed"] > 0
                and cleanup_cfg.get("enabled", False)
                and result["cleanup_all_chunks"]
            ):
                logger.info(
                    f"Đang retry {result['cleanup_failed']} chunks AI Cleanup failed..."
                )
                api_keys = cleanup_cfg.get("api_keys", [])
                if not api_keys:
                    api_keys = ocr_cfg.get("_root_api_keys", [])
                model_name = cleanup_cfg.get("model", "gemini-2.5-flash")
                prompt = """Bạn là một AI chuyên dọn dẹp văn bản OCR/scan. Nhiệm vụ:
1. Loại bỏ header/footer lặp lại ở mỗi trang
2. Loại bỏ các ký tự rác, vệt đen vô nghĩa từ quá trình scan
3. Loại bỏ số trang, watermark
4. Giữ nguyên nội dung chính của văn bản
5. Chuẩn hóa khoảng trắng thừa
6. Giữ nguyên định dạng đoạn văn

Trả về chỉ văn bản đã được dọn dẹp, không giải thích thêm.

Văn bản cần dọn dẹp:
"""
                retry_results, still_failed = _retry_failed_chunks_cleanup(
                    result["cleanup_failed_indices"],
                    result["cleanup_all_chunks"],
                    api_keys,
                    model_name,
                    prompt,
                    ocr_cfg,
                )

                all_chunks = list(result["cleanup_all_chunks"])
                for idx, retry_text in retry_results.items():
                    if idx < len(all_chunks):
                        all_chunks[idx] = retry_text
                updated_text = "\n\n".join(all_chunks)
                logger.info(
                    f"AI Cleanup Retry: {result['cleanup_failed'] - len(still_failed)}/{result['cleanup_failed']} chunks retry thành công."
                )

            if result["spell_check_failed"] > 0 and spell_check_cfg.get(
                "enabled", False
            ):
                logger.info(
                    f"Đang retry {result['spell_check_failed']} chunks AI Spell Check failed..."
                )
                api_keys = spell_check_cfg.get("api_keys", [])
                if not api_keys:
                    api_keys = ocr_cfg.get("_root_api_keys", [])
                model_name = spell_check_cfg.get("model", "gemini-2.5-flash")
                prompt = """Bạn là một AI chuyên soát lỗi chính tả và phục hồi cấu trúc văn bản OCR. Nhiệm vụ chính của bạn là PHÂN TÍCH NGỮ CẢNH và QUYẾT ĐỊNH THÔNG MINH.

=== NHIỆM VỤ CHÍNH: PHÂN TÍCH VÀ PHỤC HỒI CÂU BỊ NGẮT (Ưu tiên cao nhất) ===

Bạn cần ĐỌC KỸ NỘI DUNG và PHÂN TÍCH để phân biệt:

A. CÂU BỊ NGẮT DO CONVERT PDF → TXT (CẦN NỐI LẠI):
   - Đọc ngữ cảnh: Nếu dòng trước chưa hoàn thành ý và dòng sau tiếp nối ý đó → nối lại
   - Ví dụ: 
     * "Our client is also the owner of Vietnam Trade Mark Registration No. 315843 for "MICROBAN"
       in Class 5 covering..." 
     → Phân tích: "in Class 5" tiếp nối câu trước → NỐI LẠI thành một câu
   
   - Dấu hiệu cần nối:
     * Dòng trước không kết thúc bằng dấu câu (. ! ?) HOẶC kết thúc bằng dấu phẩy, hai chấm
     * Dòng sau bắt đầu bằng chữ thường (tiếp nối câu trước)
     * Nội dung dòng sau về mặt ngữ pháp và ngữ nghĩa là phần tiếp theo của câu trước
     * Đọc toàn bộ ngữ cảnh để hiểu rõ mối quan hệ

B. NGẮT PARAGRAPH CÓ CHỦ ĐÍCH (KHÔNG NỐI):
   - Đọc ngữ cảnh: Nếu dòng sau là ý mới, chủ đề mới, hoặc đoạn văn mới → KHÔNG nối
   - Ví dụ:
     * "...attached as Exhibit 1.
       
       Khách hàng của chúng tôi là chủ sở hữu..."
     → Phân tích: Đây là đoạn mới (chuyển từ tiếng Anh sang tiếng Việt) → KHÔNG NỐI
   
   - Dấu hiệu KHÔNG nối:
     * Dòng trước kết thúc bằng dấu chấm (. ! ?) và dòng sau bắt đầu bằng chữ hoa
     * Dòng sau là câu đầu tiên của một đoạn mới (ý tưởng mới, chủ đề mới)
     * Có sự thay đổi rõ ràng về ngữ cảnh (ví dụ: chuyển từ phần này sang phần khác)
     * Đọc toàn bộ ngữ cảnh để xác định đây là ngắt đoạn có chủ đích

QUY TRÌNH PHÂN TÍCH:
1. ĐỌC toàn bộ văn bản để hiểu cấu trúc và ngữ cảnh
2. PHÂN TÍCH từng vị trí ngắt dòng:
   - Xem xét nội dung trước và sau dòng ngắt
   - Đánh giá mối quan hệ ngữ pháp và ngữ nghĩa
   - Xác định đây là câu bị ngắt hay ngắt đoạn có chủ đích
3. QUYẾT ĐỊNH:
   - Nếu là câu bị ngắt → NỐI lại (thay line break bằng space)
   - Nếu là ngắt đoạn có chủ đích → GIỮ NGUYÊN (có thể thêm dòng trống nếu cần)
4. ÁP DỤNG nhất quán cho toàn bộ văn bản

=== CÁC NHIỆM VỤ KHÁC ===

1. SOÁT LỖI CHÍNH TẢ:
   - Sửa các lỗi chính tả do OCR (ví dụ: "Kíng" → "Kính", "hang" → "hàng")
   - Sửa các lỗi chính tả thông thường
   - KHÔNG thay đổi từ ngữ chuyên ngành, tên riêng, địa danh
   - KHÔNG thay đổi số liệu, ngày tháng, địa chỉ

2. PHỤC HỒI CẤU TRÚC PARAGRAPH:
   - Sau khi đã nối các câu bị ngắt, xác định các ngắt đoạn hợp lý
   - Mỗi đoạn văn nên có một ý chính hoàn chỉnh
   - Giữ nguyên các dòng trống giữa các đoạn đã được xác định là có chủ đích
   - Đảm bảo các câu trong một đoạn có liên quan với nhau

3. BẢO VỆ TOÀN VẸN NỘI DUNG:
   - TUYỆT ĐỐI KHÔNG thay đổi ý nghĩa của văn bản
   - KHÔNG thêm, bớt, hoặc diễn giải lại nội dung
   - KHÔNG thay đổi thứ tự từ trong câu (chỉ nối lại khi cần)
   - GIỮ NGUYÊN định dạng đặc biệt (bullet points, numbered lists, bảng)
   - GIỮ NGUYÊN các từ viết hoa nếu chúng là tên riêng, thuật ngữ

4. ĐỊNH DẠNG:
   - Giữ nguyên định dạng văn bản song ngữ (nếu có)
   - Giữ nguyên các dấu câu quan trọng
   - Chuẩn hóa khoảng trắng thừa giữa các từ (nhưng không thay đổi paragraph breaks hợp lý)
   - Đảm bảo mỗi câu kết thúc bằng dấu câu thích hợp

=== NGUYÊN TẮC QUAN TRỌNG ===

- SỬ DỤNG SỨC MẠNH PHÂN TÍCH NGỮ CẢNH: Đọc và hiểu nội dung, không chỉ dựa vào quy tắc cú pháp
- QUYẾT ĐỊNH THÔNG MINH: Mỗi quyết định nối hay không nối phải dựa trên phân tích ngữ cảnh cụ thể
- NHẤT QUÁN: Áp dụng cùng một tiêu chuẩn phân tích cho toàn bộ văn bản
- BẢO TOÀN Ý NGHĨA: Chỉ điều chỉnh cấu trúc, KHÔNG thay đổi nội dung hoặc ý nghĩa

Trả về chỉ văn bản đã được soát và phục hồi, không giải thích thêm.

Văn bản cần phân tích và xử lý:
"""

                spell_check_chunks = [
                    updated_text[i : i + spell_check_cfg.get("chunk_size", 10000)]
                    for i in range(
                        0, len(updated_text), spell_check_cfg.get("chunk_size", 10000)
                    )
                ]
                retry_results, still_failed = _retry_failed_chunks_spell_check(
                    result["spell_check_failed_indices"],
                    spell_check_chunks,
                    api_keys,
                    model_name,
                    prompt,
                    ocr_cfg,
                )

                spell_check_chunks_list = list(spell_check_chunks)
                for idx, retry_text in retry_results.items():
                    if idx < len(spell_check_chunks_list):
                        spell_check_chunks_list[idx] = retry_text
                updated_text = "\n\n".join(spell_check_chunks_list)
                logger.info(
                    f"AI Spell Check Retry: {result['spell_check_failed'] - len(still_failed)}/{result['spell_check_failed']} chunks retry thành công."
                )

            with open(output_path, "w", encoding="utf-8") as f:
                f.write(updated_text)
            logger.info(f"OCR: Đã lưu text đã được retry vào: {output_path}")
        elif user_choice == "save":
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(result_text)
            logger.info(f"OCR: Đã lưu vào: {output_path}")
        elif user_choice == "exit":
            logger.info("Thoát không lưu.")
        else:
            # Auto-save
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(result_text)
            logger.info(f"OCR: Đã tự động lưu vào: {output_path}")

    logger.info("Hoàn tất OCR pipeline.")
