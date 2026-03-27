# -*- coding: utf-8 -*-
from __future__ import annotations

import asyncio
import logging
import os
from pathlib import Path
from typing import Any, Dict, List, Optional

from .exceptions import DOCXProcessorError

# Thư viện xử lý Word sẽ được lazy import hoặc mong đợi từ ocr_reader truyền vào
# Tuy nhiên để module hoạt động độc lập, ta sẽ xử lý refs cẩn thận
logger = logging.getLogger("NovelTranslator")

# Placeholder cho các định dạng từ python-docx
# Sẽ được cập nhật động từ ocr_reader hoặc import local
Document = None
Inches = None
Pt = None
WD_PARAGRAPH_ALIGNMENT = None
Image = None

def _get_docx_refs() -> None:
    """Lấy các reference cần thiết cho docx từ ocr_reader hoặc import."""
    global Document, Inches, Pt, WD_PARAGRAPH_ALIGNMENT, Image
    try:
        from docx import Document as Doc
        from docx.enum.paragraph import WD_PARAGRAPH_ALIGNMENT as Align
        from docx.shared import Inches as Inc
        from docx.shared import Pt as P
        from PIL import Image as PILImage
        Document, Inches, Pt, WD_PARAGRAPH_ALIGNMENT, Image = Doc, Inc, P, Align, PILImage
    except ImportError:
        pass

def extract_format_hints(para: Any, para_index: int, total_paragraphs: int) -> dict:
    """Extract format hints chi tiết từ paragraph."""
    _get_docx_refs()
    hints = {
        "style": para.style.name if para.style else "Normal",
        "font_size": None,
        "is_bold": False,
        "is_italic": False,
        "alignment": "left",
        "position_hint": "middle",
    }

    # Get alignment
    if para.alignment is not None and WD_PARAGRAPH_ALIGNMENT is not None:
        if para.alignment == WD_PARAGRAPH_ALIGNMENT.LEFT:
            hints["alignment"] = "left"
        elif para.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            hints["alignment"] = "center"
        elif para.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            hints["alignment"] = "right"
        elif para.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
            hints["alignment"] = "justify"

    # Get font info từ runs
    if para.runs:
        first_run = para.runs[0]
        if first_run.font.size:
            hints["font_size"] = first_run.font.size.pt
        hints["is_bold"] = first_run.font.bold is True
        hints["is_italic"] = first_run.font.italic is True

    # Estimate position
    if total_paragraphs > 0:
        position_ratio = para_index / total_paragraphs
        if position_ratio < 0.2:
            hints["position_hint"] = "top"
        elif position_ratio > 0.8:
            hints["position_hint"] = "bottom"
        else:
            hints["position_hint"] = "middle"

    return hints

def is_in_table(para) -> bool:
    """Check nếu paragraph nằm trong table."""
    try:
        parent = para._element.getparent()
        if parent is not None:
            return parent.tag.endswith("tbl")
    except Exception:
        pass
    return False

def extract_images_from_paragraph(para: Any) -> List[dict]:
    """Extract images từ paragraph."""
    _get_docx_refs()
    images = []
    try:
        for run_idx, run in enumerate(para.runs):
            blips = run._element.xpath(".//a:blip")
            if blips:
                try:
                    blip = blips[0]
                    rId = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    if rId and hasattr(para.part, "rels") and rId in para.part.rels:
                        image_part_rel = para.part.rels[rId]
                        image_blob = image_part_rel.target_part.blob
                        width_inches, height_inches = 4.0, 3.0
                        try:
                            if Image is not None:
                                import io
                                img = Image.open(io.BytesIO(image_blob))
                                width_inches = img.width / 96.0
                                height_inches = img.height / 96.0
                                img.close()
                        except Exception:
                            pass
                        images.append({
                            "run_index": run_idx,
                            "image_data": image_blob,
                            "width": width_inches,
                            "height": height_inches,
                            "run": run,
                        })
                except Exception as e:
                    logger.debug(f"Không thể extract image từ run {run_idx}: {e}")
                    continue
    except Exception as e:
        logger.debug(f"Lỗi khi extract images từ paragraph: {e}")
    return images

def extract_paragraphs_with_hints(docx_path: str) -> List[dict]:
    """Extract paragraphs từ DOCX với format hints chi tiết và images."""
    _get_docx_refs()
    if Document is None:
        raise DOCXProcessorError("python-docx chưa được cài đặt.")
    doc = Document(docx_path)
    paragraphs_data = []
    all_paragraphs = [para for para in doc.paragraphs if not is_in_table(para)]
    total_paragraphs = len(all_paragraphs)
    for para_idx, para in enumerate(all_paragraphs):
        hints = extract_format_hints(para, para_idx, total_paragraphs)
        images = extract_images_from_paragraph(para)
        paragraphs_data.append({
            "index": len(paragraphs_data),
            "text": para.text,
            "hints": hints,
            "images": images,
            "has_images": len(images) > 0,
            "para_object": para,
        })
    logger.info(f"📝 Đã extract {len(paragraphs_data)} paragraphs từ DOCX (đã skip tables)")
    return paragraphs_data

def batch_small_paragraphs(paragraphs_data: List[dict], min_chars: int = 50) -> List[dict]:
    """Batch các paragraphs nhỏ lại với nhau để tối ưu API calls."""
    batched, current_batch = [], []
    for para in paragraphs_data:
        if len(para["text"]) < min_chars and not para.get("has_images", False):
            current_batch.append(para)
        else:
            if current_batch:
                batched.append({
                    "type": "batch",
                    "text": "\n\n".join([p["text"] for p in current_batch]),
                    "original_indices": [p["index"] for p in current_batch],
                    "para_objects": [p["para_object"] for p in current_batch],
                    "images_list": [p["images"] for p in current_batch],
                    "hints": current_batch[0]["hints"],
                })
                current_batch = []
            batched.append({"type": "single", **para})
    if current_batch:
        batched.append({
            "type": "batch",
            "text": "\n\n".join([p["text"] for p in current_batch]),
            "original_indices": [p["index"] for p in current_batch],
            "para_objects": [p["para_object"] for p in current_batch],
            "images_list": [p["images"] for p in current_batch],
            "hints": current_batch[0]["hints"],
        })
    if batched:
        batch_count = sum(1 for b in batched if b["type"] == "batch")
        single_count = sum(1 for b in batched if b["type"] == "single")
        logger.info(f"📦 Đã batch: {batch_count} batches, {single_count} single paragraphs")
    return batched

def build_cleanup_prompt_with_hints(text: str, hints: dict) -> str:
    """Build cleanup prompt với format hints chi tiết."""
    hint_descriptions = []
    style = hints.get("style", "")
    if style.startswith("Heading"):
        hint_descriptions.append(f"Style: {style} (có thể là header/title)")
    elif style == "Normal":
        hint_descriptions.append("Style: Normal (nội dung chính)")
    
    font_size = hints.get("font_size")
    if font_size:
        if font_size < 10:
            hint_descriptions.append(f"Font rất nhỏ ({font_size}pt) - có thể là footer/page number")
        elif font_size > 14:
            hint_descriptions.append(f"Font lớn ({font_size}pt) - có thể là header/title")
    
    if hints.get("is_bold"):
        hint_descriptions.append("Bold - có thể là header/title")
    
    position = hints.get("position_hint", "middle")
    if position == "top":
        hint_descriptions.append("Vị trí: Đầu trang - có thể là header")
    elif position == "bottom":
        hint_descriptions.append("Vị trí: Cuối trang - có thể là footer/page number")
    
    alignment = hints.get("alignment", "left")
    if alignment == "center":
        hint_descriptions.append("Căn giữa - có thể là title/header")

    return f"""Bạn là AI chuyên dọn dẹp văn bản OCR/scan.

THÔNG TIN FORMATTING:
{chr(10).join("- " + d for d in hint_descriptions) if hint_descriptions else "- Không có thông tin đặc biệt"}

Dựa trên formatting này, xác định:
- Nếu là header/footer/page number → XÓA
- Nếu là nội dung chính → GIỮ LẠI và cleanup noise

Nhiệm vụ:
1. Loại bỏ header/footer lặp lại ở mỗi trang
2. Loại bỏ số trang, watermark
3. Loại bỏ các ký tự rác, vệt đen vô nghĩa từ quá trình scan
4. Chuẩn hóa khoảng trắng thừa
5. Giữ nguyên nội dung chính của văn bản
6. Giữ nguyên định dạng đoạn văn

Trả về chỉ văn bản đã được dọn dẹp, không giải thích thêm.

Văn bản cần dọn dẹp:
{text}"""

# Worker ID dùng cho docx cleanup/spell check khi dùng key_manager chung (trùng với ai_processor)
_WORKER_ID_DOCX_CLEANUP = 998
_WORKER_ID_DOCX_SPELL = 997


async def cleanup_paragraph_with_hints(
    para_data: dict, ocr_cfg: dict, key_manager: Any = None
) -> dict:
    """Cleanup một paragraph/batch với format hints. Nếu key_manager có thì dùng chung pool/cooldown."""
    from .ocr_reader import _build_safety_settings, _cleanup_chunk_async

    cleanup_cfg = ocr_cfg.get("ai_cleanup", {})
    if not cleanup_cfg.get("enabled", False):
        return {"cleaned_text": para_data["text"], "should_merge_with_next": False}

    api_keys = cleanup_cfg.get("api_keys", []) or ocr_cfg.get("_root_api_keys", [])
    if key_manager is not None:
        api_keys = getattr(key_manager, "api_keys", api_keys) or api_keys
    if not api_keys:
        logger.warning("Không có API keys cho cleanup, bỏ qua")
        return {"cleaned_text": para_data["text"], "should_merge_with_next": False}

    model_name = cleanup_cfg.get("model", "gemini-2.0-flash")
    timeout_s = cleanup_cfg.get("timeout", 60.0)
    safety_level = cleanup_cfg.get("safety_level") or ocr_cfg.get("safety_level", "BLOCK_ONLY_HIGH")
    safety_settings = _build_safety_settings(safety_level)

    text = para_data["text"]
    hints = para_data.get("hints", {})
    prompt = build_cleanup_prompt_with_hints(text, hints)

    # [v9.1] get_available_key is async
    key = await key_manager.get_available_key() if key_manager else None
    if key is None and not key_manager:
        key = api_keys[0] if api_keys else None
    if key is None:
        logger.warning("Không lấy được key cho cleanup paragraph.")
        return {"cleaned_text": text, "should_merge_with_next": False}

    failed = False
    err_type, err_msg = "generation_error", ""
    try:
        cleaned_text = await _cleanup_chunk_async(
            text, key, model_name, prompt, 0, 1, timeout_s, safety_settings
        )
        return {"cleaned_text": cleaned_text, "should_merge_with_next": False}
    except Exception as e:
        failed = True
        logger.warning(f"Cleanup paragraph thất bại: {e}")
        err_type = (
            key_manager.handle_exception(key, e)
            if (key_manager and hasattr(key_manager, "handle_exception"))
            else "generation_error"
        )
        err_msg = str(e)
        return {"cleaned_text": text, "should_merge_with_next": False}
    finally:
        if key_manager and key:
            await key_manager.return_key(
                _WORKER_ID_DOCX_CLEANUP,
                key,
                is_error=failed,
                error_type=err_type,
                error_message=err_msg,
            )

def spell_check_paragraph(
    para_data: dict, ocr_cfg: dict, key_manager: Any = None
) -> str:
    """Spell check một paragraph/batch. Nếu key_manager có thì dùng chung pool/cooldown (khi có gọi AI thật)."""
    spell_check_cfg = ocr_cfg.get("ai_spell_check", {})
    if not spell_check_cfg.get("enabled", False):
        return para_data.get("cleaned_text", para_data["text"])

    api_keys = spell_check_cfg.get("api_keys", []) or ocr_cfg.get("_root_api_keys", [])
    if key_manager is not None:
        api_keys = getattr(key_manager, "api_keys", api_keys) or api_keys
    if not api_keys:
        logger.warning("Không có API keys cho spell check, bỏ qua")
        return para_data.get("cleaned_text", para_data["text"])

    text = para_data.get("cleaned_text", para_data["text"])
    # Placeholder: chưa gọi AI spell check thật. Khi bật sẽ get key, gọi _cleanup_chunk_async với spell prompt, return_key.
    return text

def update_paragraph_in_place(para: Any, new_text: str) -> None:
    """Update paragraph text nhưng giữ nguyên formatting từ original runs."""
    if not para.runs:
        para.add_run(new_text)
        return
    first_run = para.runs[0]
    font_format = {
        "bold": first_run.font.bold,
        "italic": first_run.font.italic,
        "size": first_run.font.size,
        "name": first_run.font.name,
    }
    para.clear()
    run = para.add_run(new_text)
    if font_format["bold"]: run.font.bold = True
    if font_format["italic"]: run.font.italic = True
    if font_format["size"]: run.font.size = font_format["size"]
    if font_format["name"]: run.font.name = font_format["name"]

def re_insert_images_to_paragraph(para: Any, images: List[Dict[str, Any]]) -> None:
    """Re-insert images vào paragraph sau khi process text."""
    _get_docx_refs()
    import io
    if not images: return
    for img_info in images:
        try:
            image_bytes = img_info.get("image_data")
            if not image_bytes or len(image_bytes) < 10: continue
            width_inches = img_info.get("width", 4.0)
            run = para.add_run()
            img_stream = io.BytesIO(image_bytes)
            img_stream.seek(0)
            try:
                max_width = 6.0
                run.add_picture(img_stream, width=Inches(min(width_inches, max_width)))
            except Exception as pic_error:
                logger.warning(f"Không thể re-insert image vào paragraph: {pic_error}")
        except Exception as e:
            logger.warning(f"Lỗi khi re-insert image: {e}")

def split_batched_result(batched_text: str, original_count: int) -> List[str]:
    """Split batched result về số paragraphs ban đầu (estimate)."""
    if not batched_text or not batched_text.strip():
        return [""] * original_count
    paragraphs = [p.strip() for p in batched_text.split("\n\n") if p.strip()]
    if len(paragraphs) == original_count:
        return paragraphs
    elif len(paragraphs) < original_count:
        paragraphs.extend([""] * (original_count - len(paragraphs)))
        return paragraphs
    else:
        if original_count <= 0: return paragraphs
        result = paragraphs[: original_count - 1]
        remaining = "\n\n".join(paragraphs[original_count - 1 :])
        result.append(remaining)
        return result

def update_docx_with_processed_text(docx_path: str, processed_paragraphs: List[dict], ocr_cfg: dict) -> str:
    """Update DOCX với processed text, giữ nguyên formatting và images."""
    _get_docx_refs()
    if Document is None: raise DOCXProcessorError("python-docx chưa được cài đặt.")
    logger.info(f"🔄 Đang update DOCX với processed text: {docx_path}")
    doc = Document(docx_path)
    para_idx = 0
    for processed in processed_paragraphs:
        if processed["type"] == "batch":
            cleaned_text = processed.get("cleaned_text", processed["text"])
            spell_checked_text = processed.get("spell_checked_text", cleaned_text)
            para_objects, images_list = processed["para_objects"], processed.get("images_list", [])
            split_paragraphs = split_batched_result(spell_checked_text, len(para_objects))
            for i, para_obj in enumerate(para_objects):
                para_text = split_paragraphs[i] if i < len(split_paragraphs) else ""
                if para_text.strip():
                    update_paragraph_in_place(para_obj, para_text)
                if i < len(images_list):
                    re_insert_images_to_paragraph(para_obj, images_list[i])
        else:
            para_obj = processed["para_object"]
            cleaned_text = processed.get("cleaned_text", processed["text"])
            spell_checked_text = processed.get("spell_checked_text", cleaned_text)
            images = processed.get("images", [])
            if spell_checked_text.strip():
                update_paragraph_in_place(para_obj, spell_checked_text)
                re_insert_images_to_paragraph(para_obj, images)
            else:
                para_obj.clear()
        
        # Merge logic (simplified)
        para_idx += 1
    doc.save(docx_path)
    logger.info(f"✅ Đã update DOCX: {docx_path}")
    return docx_path

def convert_docx_to_epub(docx_path: str, epub_path: str, ocr_cfg: dict) -> str:
    """Convert DOCX → EPUB using pypandoc."""
    try:
        import pypandoc
    except ImportError:
        raise DOCXProcessorError("pypandoc chưa được cài đặt.")
    logger.info(f"📚 Đang convert DOCX → EPUB: {docx_path}")
    try:
        pypandoc.convert_file(docx_path, "epub", outputfile=epub_path, extra_args=["--standalone"])
        return epub_path
    except Exception as e:
        logger.error(f"❌ Lỗi khi convert DOCX → EPUB: {e}")
        raise

def create_docx_from_processed_text(pages_data: List[dict], output_path: str, ocr_cfg: dict, pdf_path: str = "") -> str:
    """Tạo DOCX từ logic extract và process."""
    _get_docx_refs()
    if Document is None: raise DOCXProcessorError("python-docx chưa được cài đặt.")
    import io

    from tqdm import tqdm

    doc = Document()
    show_progress = bool(ocr_cfg.get("show_progress", True))
    
    # Insert images
    all_images = []
    for page_info in pages_data:
        all_images.extend([(page_info["page_num"], img_info) for img_info in page_info.get("images", [])])
    
    if all_images:
        for page_num, img_info in (tqdm(all_images, desc="Chèn images") if show_progress else all_images):
            try:
                image_bytes = img_info.get("data")
                if not image_bytes or len(image_bytes) < 10: continue
                para = doc.add_paragraph()
                run = para.add_run()
                img_stream = io.BytesIO(image_bytes)
                img_width = img_info.get("width", 500)
                width_inches = min(6.0, img_width / 96.0)
                run.add_picture(img_stream, width=Inches(width_inches))
            except Exception: continue

    # Insert text
    all_text = "\n\n".join([p["text"] for p in pages_data])
    for para_text in all_text.split("\n\n"):
        if para_text.strip():
            doc.add_paragraph(para_text.strip())

    doc.save(output_path)
    return output_path

def _create_html_from_items(
    all_items_with_position: List[dict], output_path: str
) -> str:
    """
    Tạo file HTML từ all_items_with_position (text + images).
    Images được embed dưới dạng base64 để không cần temp files.

    Args:
        all_items_with_position: List các items (text hoặc image) đã được sort theo (Y, X)
        output_path: Đường dẫn file DOCX output (để tạo HTML temp file cùng folder)

    Returns:
        str: Đường dẫn file HTML đã tạo
    """
    import base64
    import html

    html_path = output_path.replace(".docx", "_temp.html")
    if html_path == output_path:  # Nếu không phải .docx
        html_path = output_path + "_temp.html"

    logger.info(f"Tao HTML trung gian: {html_path}")

    html_parts = ['<!DOCTYPE html>\n<html>\n<head>\n<meta charset="UTF-8">\n']
    html_parts.append("<style>\n")
    html_parts.append(
        "body { font-family: Arial, sans-serif; margin: 20px; line-height: 1.6; }\n"
    )
    html_parts.append("p { margin-bottom: 6pt; }\n")
    html_parts.append(
        "img { max-width: 6in; height: auto; margin: 6pt 0; display: block; }\n"
    )
    html_parts.append("</style>\n</head>\n<body>\n")

    images_count = 0
    text_count = 0

    for item in all_items_with_position:
        if item["type"] == "text":
            content = item["content"]
            # Escape HTML và thay thế line breaks
            content = html.escape(content)
            content = content.replace("\n\n", "</p><p>")
            content = content.replace("\n", "<br>")
            html_parts.append(f"<p>{content}</p>\n")
            text_count += 1
        elif item["type"] == "image":
            img_info = item["img_info"]
            image_bytes = img_info.get("data")
            if image_bytes and len(image_bytes) >= 10:
                image_ext = img_info.get("ext", "png").lower()
                if image_ext == "jpg":
                    image_ext = "jpeg"

                # Convert image bytes to base64
                try:
                    base64_data = base64.b64encode(image_bytes).decode("utf-8")
                    data_uri = f"data:image/{image_ext};base64,{base64_data}"

                    # Get image dimensions for sizing
                    img_width = img_info.get("width", 0)
                    img_height = img_info.get("height", 0)

                    if img_width > 0 and img_height > 0:
                        max_width_px = 576  # 6 inches at 96 DPI
                        if img_width > max_width_px:
                            aspect_ratio = img_height / img_width
                            display_width = max_width_px
                            display_height = int(max_width_px * aspect_ratio)
                        else:
                            display_width = img_width
                            display_height = img_height

                        html_parts.append(
                            f'<img src="{data_uri}" width="{display_width}" height="{display_height}" alt="Image from page {item["page_num"]}">\n'
                        )
                    else:
                        html_parts.append(
                            f'<img src="{data_uri}" alt="Image from page {item["page_num"]}">\n'
                        )

                    images_count += 1
                    logger.debug(
                        f"Da embed image {images_count} vao HTML (trang {item['page_num']}, size: {len(image_bytes)} bytes)"
                    )
                except Exception as e:
                    logger.warning(
                        f"Khong the convert image tu trang {item['page_num']} sang base64: {e}"
                    )
            else:
                logger.warning(
                    f"Image tu trang {item['page_num']} khong co data hop le (size: {len(image_bytes) if image_bytes else 0} bytes)"
                )

    html_parts.append("</body>\n</html>")

    html_content = "".join(html_parts)

    try:
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        logger.info(f"Da tao HTML: {text_count} paragraphs, {images_count} images")
        return html_path
    except Exception as e:
        logger.error(f"Khong the tao HTML file: {e}")
        raise


def _convert_html_to_docx_with_pandoc(
    html_path: str, output_path: str, ocr_cfg: dict
) -> bool:
    """
    Convert HTML sang DOCX bằng pandoc.

    Args:
        html_path: Đường dẫn file HTML input
        output_path: Đường dẫn file DOCX output
        ocr_cfg: Config dictionary

    Returns:
        bool: True nếu thành công, False nếu thất bại
    """
    try:
        import pypandoc
    except ImportError:
        logger.warning(
            "pypandoc chua duoc cai dat. Cai pypandoc de dung HTML intermediate workflow."
        )
        return False

    try:
        logger.info("Dang convert HTML -> DOCX bang pandoc...")
        # Pandoc options để preserve images và formatting
        extra_args = [
            "--standalone",
            "--wrap=none",  # Không wrap lines
        ]

        pypandoc.convert_file(
            html_path, "docx", outputfile=output_path, extra_args=extra_args
        )

        logger.info("Da convert HTML -> DOCX thanh cong bang pandoc")
        return True
    except Exception as e:
        logger.warning(f"Pandoc conversion that bai: {e}")
        import traceback

        logger.debug(traceback.format_exc())
        return False

