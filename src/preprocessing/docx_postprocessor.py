# -*- coding: utf-8 -*-
from __future__ import annotations

"""
DOCX Post-Processor: Xóa header/footer lặp và watermark/shapes sau khi convert PDF→DOCX bằng pdf2docx.
Heuristic an toàn: xóa các đoạn ngắn lặp với tần suất cao; xóa shapes/drawing chứa từ khóa watermark.
"""

import logging
import re
from typing import Any, Dict, List

logger = logging.getLogger("NovelTranslator")

try:
    from docx import Document  # type: ignore
except Exception:
    Document = None


def _normalize_text(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r"\s+", " ", s)
    return s


def _collect_paragraph_texts(doc) -> List[str]:
    texts: List[str] = []
    for p in doc.paragraphs:
        texts.append(_normalize_text(p.text))
    return texts


def _find_repeated_lines(
    lines: List[str], max_len: int, threshold_ratio: float
) -> List[str]:
    freq: Dict[str, int] = {}
    for t in lines:
        if not t:
            continue
        if len(t) > max_len:
            continue
        freq[t] = freq.get(t, 0) + 1
    if not freq:
        return []
    n = max(1, len(lines))
    repeated = [t for t, c in freq.items() if (c / n) >= threshold_ratio]
    return repeated


def _remove_paragraphs_matching(
    doc, patterns: List[str], case_insensitive: bool
) -> int:
    removed = 0
    if not patterns:
        return 0
    flags = re.IGNORECASE if case_insensitive else 0
    regex_list = [re.compile(re.escape(pat), flags) for pat in patterns if pat]
    if not regex_list:
        return 0
    for p in list(doc.paragraphs):
        txt = _normalize_text(p.text)
        if not txt:
            continue
        if any(rgx.search(txt) for rgx in regex_list):
            # Xóa paragraph này
            p._element.getparent().remove(p._element)
            removed += 1
    return removed


def _remove_shapes_with_keywords(
    doc, keywords: List[str], case_insensitive: bool
) -> int:
    """Xóa các paragraph chứa drawing/shapes và có text trùng từ khóa (simple heuristic)."""
    removed = 0
    if not keywords:
        return 0
    flags = re.IGNORECASE if case_insensitive else 0
    rgx = (
        re.compile("|".join(re.escape(k) for k in keywords if k), flags)
        if any(keywords)
        else None
    )
    for p in list(doc.paragraphs):
        try:
            has_drawing = bool(
                p._element.xpath(".//w:drawing") or p._element.xpath(".//v:shape")
            )
        except Exception:
            has_drawing = False
        if not has_drawing:
            continue
        txt = _normalize_text(p.text)
        if rgx is None or rgx.search(txt or ""):
            p._element.getparent().remove(p._element)
            removed += 1
    # Header/Footer shapes
    for section in getattr(doc, "sections", []):
        for part in [section.header, section.footer]:
            try:
                hdr_elm = part.part.element
                shapes = hdr_elm.xpath(".//w:drawing") + hdr_elm.xpath(".//v:shape")
                if shapes:
                    # Xóa toàn bộ đoạn chứa shapes bằng cách xóa parent w:p gần nhất
                    for shp in shapes:
                        try:
                            p = shp.getparent()
                            while p is not None and p.tag and not p.tag.endswith("p"):
                                p = p.getparent()
                            if p is not None and p.getparent() is not None:
                                p.getparent().remove(p)
                                removed += 1
                        except Exception:
                            continue
            except Exception:
                continue
    return removed


def _remove_all_headers_footers(doc) -> int:
    """Xóa toàn bộ paragraph/tables trong header & footer của tất cả sections."""
    removed = 0
    for section in getattr(doc, "sections", []):
        for part in [section.header, section.footer]:
            try:
                # Xóa paragraphs
                for p in list(part.paragraphs):
                    p._element.getparent().remove(p._element)
                    removed += 1
                # Xóa tables
                for t in list(part.tables):
                    t._element.getparent().remove(t._element)
                    removed += 1
            except Exception:
                continue
    return removed


def _remove_page_number_fields(doc) -> int:
    """Gỡ các field PAGE/NUMPAGES trong tài liệu (header/footer/body)."""
    removed = 0
    root = doc.part.element
    try:
        # fldSimple có instr trực tiếp
        fld_nodes = root.xpath(
            ".//w:fldSimple[contains(@w:instr, 'PAGE') or contains(@w:instr, 'NUMPAGES')]",
            namespaces=root.nsmap,
        )
        for fld in fld_nodes:
            p = fld.getparent()
            while p is not None and (not p.tag or not p.tag.endswith("p")):
                p = p.getparent()
            if p is not None and p.getparent() is not None:
                p.getparent().remove(p)
                removed += 1
    except Exception:
        pass
    try:
        # Trường hợp field tách rời: w:instrText chứa PAGE/NUMPAGES
        instr_nodes = root.xpath(
            ".//w:instrText[contains(., 'PAGE') or contains(., 'NUMPAGES')]",
            namespaces=root.nsmap,
        )
        for instr in instr_nodes:
            p = instr.getparent()
            while p is not None and (not p.tag or not p.tag.endswith("p")):
                p = p.getparent()
            if p is not None and p.getparent() is not None:
                p.getparent().remove(p)
                removed += 1
    except Exception:
        pass
    return removed


def _remove_body_shapes(doc) -> int:
    """Xóa mọi drawing/shape còn lại trong body như watermark ảnh/vector."""
    removed = 0
    try:
        body = doc.part.element
        shapes = body.xpath(".//w:drawing | .//v:shape", namespaces=body.nsmap)
        for shp in shapes:
            p = shp.getparent()
            while p is not None and (not p.tag or not p.tag.endswith("p")):
                p = p.getparent()
            if p is not None and p.getparent() is not None:
                p.getparent().remove(p)
                removed += 1
    except Exception:
        pass
    return removed


def postprocess_docx(docx_path: str, config: Dict[str, Any]) -> bool:
    """
    Xử lý DOCX in-place: xóa header/footer lặp và watermark shapes theo config.
    """
    if Document is None:
        logger.warning("python-docx không khả dụng, bỏ qua hậu xử lý DOCX.")
        return False
    try:
        doc = Document(docx_path)
    except Exception as e:
        logger.warning(f"Không thể mở DOCX để hậu xử lý: {e}")
        return False

    pp_cfg = (config or {}).get("ocr", {}).get("docx_postprocess", {})
    if not pp_cfg.get("enabled", True):
        return False

    removed_total = 0

    # 1) Xóa header/footer dạng đoạn lặp
    if pp_cfg.get("remove_repeated_header_footer", True):
        max_len = int(pp_cfg.get("max_line_length", 120) or 120)
        thr = float(pp_cfg.get("repetition_threshold_ratio", 0.6) or 0.6)
        lines = _collect_paragraph_texts(doc)
        repeated = _find_repeated_lines(lines, max_len, thr)
        if repeated:
            removed_total += _remove_paragraphs_matching(doc, repeated, True)

    # 2) Xóa watermark/shapes theo từ khóa
    if pp_cfg.get("remove_watermark_shapes", True):
        keywords = pp_cfg.get("watermark_keywords", []) or []
        ci = bool(pp_cfg.get("case_insensitive", True))
        removed_total += _remove_shapes_with_keywords(doc, keywords, ci)

    # 2.1) Xóa toàn bộ header/footer nếu bật cờ mạnh tay
    if pp_cfg.get("remove_all_headers_footers", False):
        removed_total += _remove_all_headers_footers(doc)

    # 2.2) Gỡ PAGE/NUMPAGES ở mọi nơi nếu bật
    if pp_cfg.get("remove_page_numbers", True):
        removed_total += _remove_page_number_fields(doc)

    # 2.3) Xóa shapes còn lại trong body nếu bật
    if pp_cfg.get("remove_body_watermark_shapes", True):
        removed_total += _remove_body_shapes(doc)

    try:
        doc.save(docx_path)
        logger.info(f"DOCX hậu xử lý: đã xóa {removed_total} phần header/footer/shapes")
        return True
    except Exception as e:
        logger.warning(f"Không thể lưu DOCX sau hậu xử lý: {e}")
        return False
